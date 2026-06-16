"""Genera el Word de propuesta curricular directamente desde los JSON, sin Streamlit."""
from __future__ import annotations

import io
import json
import sys
from datetime import date
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

BASE_DIR   = Path(__file__).resolve().parent
LOGO_PATH  = BASE_DIR / "assets" / "logo_shoa.png"
OUTPUT_DOC = BASE_DIR / "Propuesta_Curricular_SHOA.docx"

REQUIRED = {
    "prog": BASE_DIR / "progresion_analisis.json",
    "reit": BASE_DIR / "reiteracion_matriz.json",
    "prof": BASE_DIR / "profundidad_iho.json",
    "prop": BASE_DIR / "propuesta_curricular.json",
}


def load(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def generar_word(prog: dict, reit: dict, prof: dict, prop: dict) -> bytes:
    doc = DocxDocument()

    def _h(text: str, level: int = 1) -> None:
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _p(text: str) -> None:
        doc.add_paragraph(text)

    def _tabla(filas: list[list[str]]) -> None:
        if not filas:
            return
        t = doc.add_table(rows=len(filas), cols=len(filas[0]))
        t.style = "Table Grid"
        for i, fila in enumerate(filas):
            for j, val in enumerate(fila):
                cell = t.rows[i].cells[j]
                cell.text = str(val)
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

    # ── Portada ──────────────────────────────────────────────────────────────
    if LOGO_PATH.exists():
        try:
            doc.add_picture(str(LOGO_PATH), width=Inches(2.0))
        except Exception:
            pass

    par_titulo = doc.add_paragraph()
    par_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = par_titulo.add_run(
        "PROPUESTA DE REDISEÑO CURRICULAR\n"
        "Programa de Hidrografía — SHOA\n"
    )
    run_t.bold = True
    run_t.font.size = Pt(18)

    par_sub = doc.add_paragraph()
    par_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par_sub.add_run(
        f"Ajuste a 1 año y medio (3 semestres)\n"
        f"Estándares IHO S-5A\n"
        f"Generado: {date.today().strftime('%d/%m/%Y')}"
    )
    doc.add_page_break()

    # ── 1. Resumen ejecutivo ─────────────────────────────────────────────────
    _h("1. Resumen Ejecutivo")
    res = prop.get("resumen", {})
    _p(
        "Este documento presenta los resultados del análisis curricular del programa de "
        "Hidrografía de la Armada de Chile (SHOA) y propone su reestructuración a 3 semestres "
        "(1 año y medio), alineado con los estándares IHO S-5A."
    )
    _tabla([
        ["Indicador", "Valor"],
        ["Horas actuales",       str(res.get("total_horas_actuales", "—"))],
        ["Horas propuestas",     str(res.get("total_horas_propuestas", "—"))],
        ["Reducción de horas",   str(res.get("reduccion_horas", "—"))],
        ["Reducción porcentual", f"{res.get('reduccion_porcentaje', 0):.1f}%"],
        ["Cobertura IHO mantenida", f"{res.get('cobertura_iho_mantenida_porcentaje', 0):.1f}%"],
        ["Cabe en 1.5 años",    "Sí" if res.get("curriculo_cabe_en_1_5_anios") else "No"],
    ])

    # ── 2. Análisis de progresión ────────────────────────────────────────────
    _h("2. Análisis de Progresión")
    r_prog = prog.get("resumen", {})
    _p(
        f"Asignaturas: {r_prog.get('total_asignaturas', 0)} | "
        f"Dependencias: {r_prog.get('total_dependencias', 0)} | "
        f"Problemas detectados: {r_prog.get('total_problemas', 0)}"
    )
    problemas_prog = prog.get("problemas_progresion", [])
    if problemas_prog:
        _h("2.1 Problemas de Progresión", level=2)
        filas_pb = [["#", "Tipo", "Asignatura", "Sem.", "Descripción"]]
        for i, pb in enumerate(problemas_prog, 1):
            desc = pb.get("descripcion", "")
            filas_pb.append([
                str(i),
                pb.get("tipo", "").replace("_", " ").title(),
                pb.get("asignatura", "—"),
                str(pb.get("semestre_asignatura", "—")),
                desc[:110] + ("…" if len(desc) > 110 else ""),
            ])
        _tabla(filas_pb)

    sin_prereq = prog.get("asignaturas_sin_prerrequisitos_claros", [])
    if sin_prereq:
        _h("2.2 Asignaturas sin Prerrequisitos Claros", level=2)
        _tabla(
            [["Código", "Nombre", "Nivel IHO"]]
            + [[a["codigo"], a.get("nombre", ""),
                str(a.get("nivel_iho_predominante", ""))] for a in sin_prereq]
        )

    # ── 3. Matriz de reiteración ─────────────────────────────────────────────
    _h("3. Matriz de Reiteración de Contenidos")
    r_reit = reit.get("resumen", {})
    _p(
        f"Pares de contenidos similares (umbral ≥{r_reit.get('threshold_fuzzy', 75)}%): "
        f"{r_reit.get('total_reiteraciones_detectadas', 0)} | "
        f"Innecesarias: {r_reit.get('total_innecesarias', 0)} | "
        f"Horas recuperables: {r_reit.get('horas_recuperables_estimadas', 0):.1f} h"
    )
    innecesarias = [r for r in reit.get("reiteraciones", [])
                    if r["tipo_reiteracion"] == "INNECESARIA"]
    if innecesarias:
        _h("3.1 Reiteraciones Innecesarias (muestra)", level=2)
        filas_ri = [["Contenido (recortado)", "Asig. 1", "Asig. 2", "Similitud", "Horas dup."]]
        for r in innecesarias[:20]:
            filas_ri.append([
                r["contenido"][:80],
                r["asignatura_1"], r["asignatura_2"],
                f"{r['similitud']}%", str(r["horas_duplicadas_estimadas"]),
            ])
        _tabla(filas_ri)

    # ── 4. Profundidad IHO ───────────────────────────────────────────────────
    _h("4. Análisis de Profundidad IHO S-5A")
    r_prof = prof.get("resumen", {})
    _p(
        f"Tópicos analizados: {r_prof.get('total_topicos_analizados', 0)} | "
        f"Sobredimensionados: {r_prof.get('sobredimensionados', 0)} | "
        f"Alineados: {r_prof.get('alineados', 0)} | "
        f"Subdimensionados: {r_prof.get('subdimensionados', 0)}"
    )
    por_asig = prof.get("por_asignatura", {})
    if por_asig:
        _h("4.1 Resumen por Asignatura", level=2)
        _tabla(
            [["Asignatura", "Sobredim.", "Alineados", "Subdim."]]
            + [[cod,
                str(v.get("SOBREDIMENSIONADO", 0)),
                str(v.get("ALINEADO", 0)),
                str(v.get("SUBDIMENSIONADO", 0))]
               for cod, v in sorted(por_asig.items())]
        )

    # ── 5. Propuesta de nueva estructura ─────────────────────────────────────
    _h("5. Propuesta de Nueva Estructura Curricular")
    _p(
        "El currículo se reorganiza en 3 semestres (1 año y medio) siguiendo la progresión IHO: "
        "Fundamentos (Sem. 1) → Operaciones (Sem. 2) → Procesamiento y Práctica (Sem. 3). "
        "Se presentan dos versiones según el tratamiento del Proyecto Hidrográfico (PH401)."
    )

    # ── 5a. Tabla de asignaturas individuales ────────────────────────────────
    _h("5.1 Tabla de Asignaturas con Acción Propuesta", level=2)
    prop_asigs = prop.get("propuesta_asignaturas", [])
    if prop_asigs:
        filas_prop = [[
            "Asignatura", "Nombre", "Acción",
            "T act.", "P act.", "AE act.", "Total act.",
            "T prop.", "P prop.", "AE prop.", "Total prop.",
            "Sem. act.", "Sem. prop.", "Justificación",
        ]]
        for pa in prop_asigs:
            ha = pa["horas_actuales"]
            hp = pa["horas_propuestas"]
            filas_prop.append([
                pa["asignatura"], pa.get("nombre_asignatura", "")[:30], pa["accion"],
                str(ha["T"]), str(ha["P"]), str(ha["SG"]), str(ha["total"]),
                str(hp["T"]), str(hp["P"]), str(hp["SG"]), str(hp["total"]),
                str(pa["semestre_actual"]), str(pa["semestre_propuesto"]),
                pa["justificacion"][:90],
            ])
        _tabla(filas_prop)

    # Helper para formatear tabla de malla
    def _tabla_malla(malla: list[dict], titulo: str, descripcion: str) -> None:
        _h(titulo, level=2)
        _p(descripcion)
        from collections import defaultdict
        sd: dict = defaultdict(list)
        for s in malla:
            sd[s["semestre"]].append(s)
        for sem in sorted(sd):
            slots   = sd[sem]
            h_total = sum(s["horas_propuestas"]["total"] for s in slots)
            h_lect  = sum(s["horas_propuestas"]["total"] for s in slots if "PH401" not in s["codigo"])
            h_campo = h_total - h_lect
            if h_campo:
                detalle = f"({h_lect} h lectivas + {h_campo} h práctica de campo)"
            else:
                detalle = f"{h_total} h lectivas"
            _h(f"Semestre {sem} — {detalle}", level=3)
            filas_m = [["Asignatura / Fusión", "Nombre", "Acción", "T", "P", "AE", "Total", "Nota"]]
            for s in slots:
                hp = s["horas_propuestas"]
                filas_m.append([
                    s["codigo"], s["nombre"][:35], s["accion"],
                    str(hp["T"]), str(hp["P"]), str(hp["SG"]), str(hp["total"]),
                    (s.get("nota") or "")[:60],
                ])
            _tabla(filas_m)

    # ── 5b. Versión A — PH401 al final ───────────────────────────────────────
    malla_va = prop.get("malla_v_a", [])
    if malla_va:
        _tabla_malla(
            malla_va,
            "5.2 Versión A — Estructura Tradicional (PH401 al final)",
            "PH401 Proyecto Hidrográfico completo en Semestre 3. "
            "Primer ciclo académico íntegro antes del trabajo de campo. "
            "Opción recomendada para grupos con menos experiencia previa.",
        )

    # ── 5c. Versión B — PH401 distribuido ────────────────────────────────────
    malla_vb = prop.get("malla_v_b", [])
    if malla_vb:
        _tabla_malla(
            malla_vb,
            "5.3 Versión B — Estructura Flexible (PH401 distribuido)",
            "PH401 dividido: 20 % introductorio en Semestre 2 (salidas a terreno iniciales) "
            "y 80 % proyecto completo en Semestre 3. "
            "Permite exposición temprana al trabajo hidrográfico de campo. "
            "Requiere coordinación logística desde el segundo semestre.",
        )

    # ── 5d. Tabla comparativa A vs B ─────────────────────────────────────────
    _h("5.4 Comparativa Versión A vs. Versión B", level=2)
    def _sem_h(malla: list[dict], sem: int) -> tuple[int, int]:
        slots  = [s for s in malla if s["semestre"] == sem]
        total  = sum(s["horas_propuestas"]["total"] for s in slots)
        lect   = sum(s["horas_propuestas"]["total"] for s in slots if "PH401" not in s["codigo"])
        return total, lect

    def _fmt_h(malla: list[dict], sem: int) -> str:
        total, lect = _sem_h(malla, sem)
        campo = total - lect
        return f"{total} h" + (f" ({lect}h lect.+{campo}h campo)" if campo else "")

    if malla_va and malla_vb:
        _tabla([
            ["Criterio", "Versión A", "Versión B"],
            ["Horas Semestre 1", _fmt_h(malla_va, 1), _fmt_h(malla_vb, 1)],
            ["Horas Semestre 2", _fmt_h(malla_va, 2), _fmt_h(malla_vb, 2)],
            ["Horas Semestre 3", _fmt_h(malla_va, 3), _fmt_h(malla_vb, 3)],
            ["Exposición temprana al campo", "No — PH401 inicia en Sem. 3",
             "Sí — salidas desde Sem. 2"],
            ["Riesgo académico", "Bajo", "Medio (logística Sem. 2)"],
            ["Flexibilidad operativa", "Baja", "Alta"],
            ["Recomendación",
             "Estructura tradicional; prerrequisitos completamente cubiertos",
             "Exposición gradual al campo; apropiada con recursos logísticos disponibles"],
        ])

    # ── 6. Métricas de impacto ───────────────────────────────────────────────
    _h("6. Métricas de Impacto")
    comp = prop.get("comparativa_internacional", {})
    dist = res.get("distribucion_semestres_horas", {})
    _tabla([
        ["Métrica", "Valor"],
        ["Horas actuales",     str(res.get("total_horas_actuales", "—"))],
        ["Horas propuestas",   str(res.get("total_horas_propuestas", "—"))],
        ["Reducción",          f"{res.get('reduccion_porcentaje', 0):.1f}%"],
        ["Cobertura IHO S-5A", f"{res.get('cobertura_iho_mantenida_porcentaje', 0):.1f}%"],
        ["Horas semestre 1",   str(dist.get("semestre_1", "—"))],
        ["Horas semestre 2",   str(dist.get("semestre_2", "—"))],
        ["Horas semestre 3",   str(dist.get("semestre_3", "—"))],
        ["Promedio internacional (tópicos)",
         f"{comp.get('promedio_internacional_topics', 0):.0f} h"],
        ["Brecha propuesta vs. intl.",
         f"{comp.get('brecha_vs_promedio_internacional', 0):+.0f} h"],
        ["Reiteraciones innecesarias",
         str(r_reit.get("total_innecesarias", "—"))],
        ["Horas recuperadas (reiteración)",
         f"{r_reit.get('horas_recuperables_estimadas', 0):.1f} h"],
    ])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def main() -> None:
    for key, path in REQUIRED.items():
        if not path.exists():
            raise FileNotFoundError(
                f"Falta {path.name} — ejecuta primero el script correspondiente."
            )

    datos = {k: load(v) for k, v in REQUIRED.items()}

    docx_bytes = generar_word(**datos)
    OUTPUT_DOC.write_bytes(docx_bytes)

    size_kb = OUTPUT_DOC.stat().st_size // 1024
    print(f"Documento Word generado: {OUTPUT_DOC}")
    print(f"Tamaño: {size_kb} KB")

    # Verificar PH401 en la propuesta
    prop_asigs = datos["prop"].get("propuesta_asignaturas", [])
    ph = next((p for p in prop_asigs if p["asignatura"] == "PH401"), None)
    if ph:
        ha = ph["horas_actuales"]
        hp = ph["horas_propuestas"]
        print(f"\nPH401 en documento:")
        print(f"  Actuales : T={ha['T']}  P={ha['P']}  AE={ha['SG']}  Total={ha['total']}")
        print(f"  Propuestas: T={hp['T']}  P={hp['P']}  AE={hp['SG']}  Total={hp['total']}")
        print(f"  Accion: {ph['accion']}")


if __name__ == "__main__":
    main()
