from __future__ import annotations

import json
import re
import sys
import unicodedata
from pathlib import Path
from typing import Any

from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

BASE_DIR = Path(__file__).resolve().parent
PROGRAMAS_DIR = BASE_DIR / "programas"
OUTPUT_PATH = BASE_DIR / "programas_estructurados.json"

CODIGOS_ASIGNATURA = [
    "ABM101",
    "GDM301",
    "GP204",
    "HDP302",
    "HO202",
    "LR303",
    "LSM102",
    "MGG103",
    "OA203",
    "PH401",
    "PhyOce104",
    "RS201",
    "SG304",
    "WTC205",
]

CODIGO_NORMALIZADO = {re.sub(r"[^a-z0-9]", "", c.lower()): c for c in CODIGOS_ASIGNATURA}


def normalizar(texto: str) -> str:
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(ch for ch in texto if not unicodedata.combining(ch))
    texto = texto.lower()
    texto = re.sub(r"\s+", " ", texto).strip()
    return texto


def normalizar_codigo(texto: str) -> str:
    base = normalizar(texto)
    return re.sub(r"[^a-z0-9]", "", base)


def extraer_lineas_documento(doc: Document) -> list[str]:
    lineas: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lineas.append(t)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parte in celda.text.splitlines():
                    t = parte.strip()
                    if t:
                        lineas.append(t)
    return lineas


def deduplicar_preservando_orden(items: list[str]) -> list[str]:
    vistos: set[str] = set()
    salida: list[str] = []
    for item in items:
        key = normalizar(item)
        if key and key not in vistos:
            vistos.add(key)
            salida.append(item)
    return salida


def buscar_primer_numero(lines: list[str], idx: int, alcance: int = 6) -> int | None:
    for j in range(idx, min(len(lines), idx + alcance)):
        m = re.search(r"\d+(?:[.,]\d+)?", lines[j])
        if m:
            valor = float(m.group(0).replace(",", "."))
            if 0 <= valor <= 5000:
                return int(round(valor))
    return None


def extraer_codigo(path_docx: Path, lineas: list[str]) -> str | None:
    stem_norm = normalizar_codigo(path_docx.stem)
    for norm, canon in CODIGO_NORMALIZADO.items():
        if norm in stem_norm:
            return canon

    patron = re.compile(r"\b([A-Za-z]{2,10}\s?-?\d{3})\b")
    for ln in lineas[:80]:
        for bruto in patron.findall(ln):
            norm = normalizar_codigo(bruto)
            if norm in CODIGO_NORMALIZADO:
                return CODIGO_NORMALIZADO[norm]
    return None


def limpiar_nombre_asignatura(texto: str) -> str:
    texto = re.sub(r"\(\s*[A-Za-z]{2,10}\s?-?\d{3}\s*\)", "", texto).strip()
    texto = re.sub(r"\s+", " ", texto).strip(" -:")
    return texto


def extraer_nombre(path_docx: Path, lineas: list[str], codigo: str | None) -> str:
    for i, ln in enumerate(lineas):
        if "nombre del programa de asignatura" in normalizar(ln):
            for j in range(i + 1, min(len(lineas), i + 8)):
                candidato = lineas[j].strip()
                if candidato and "total de horas" not in normalizar(candidato):
                    limpio = limpiar_nombre_asignatura(candidato)
                    if limpio:
                        return limpio
    for ln in lineas[:40]:
        if re.search(r"\([A-Za-z]{2,10}\s?-?\d{3}\)", ln):
            limpio = limpiar_nombre_asignatura(ln)
            if limpio:
                return limpio
    nombre_archivo = path_docx.stem
    if codigo:
        nombre_archivo = re.sub(re.escape(codigo), "", nombre_archivo, flags=re.IGNORECASE)
    nombre_archivo = re.sub(r"[_\-]+", " ", nombre_archivo).strip()
    return nombre_archivo


def extraer_horas(lineas: list[str]) -> dict[str, int]:
    horas: dict[str, int | None] = {"T": None, "P": None, "SG": None, "total": None}
    patrones = {
        "T": ["horas teoricas", "hrs teoricas", "teoricas"],
        "P": ["horas practicas", "hrs practicas", "practicas", "tutoriales"],
        "SG": ["horas autoestudio", "autoestudio", "auto estudio", "sg"],
        "total": ["total de horas", "horas totales", "total horas"],
    }

    for i, ln in enumerate(lineas):
        n = normalizar(ln)
        for campo, pats in patrones.items():
            if horas[campo] is None and any(p in n for p in pats):
                horas[campo] = buscar_primer_numero(lineas, i)

    if horas["T"] is None or horas["P"] is None or horas["SG"] is None:
        idx_total = next(
            (i for i, ln in enumerate(lineas) if "total de horas" in normalizar(ln)),
            None,
        )
        if idx_total is not None:
            nums: list[int] = []
            for j in range(idx_total, min(len(lineas), idx_total + 30)):
                m = re.fullmatch(r"\d{1,4}", lineas[j].strip())
                if m:
                    nums.append(int(m.group(0)))
            if len(nums) >= 3:
                horas["T"] = horas["T"] if horas["T"] is not None else nums[0]
                horas["P"] = horas["P"] if horas["P"] is not None else nums[1]
                horas["SG"] = horas["SG"] if horas["SG"] is not None else nums[2]
            if len(nums) >= 4 and horas["total"] is None:
                horas["total"] = nums[3]

    t = int(horas["T"] or 0)
    p = int(horas["P"] or 0)
    sg = int(horas["SG"] or 0)
    total = int(horas["total"] or (t + p + sg))
    return {"T": t, "P": p, "SG": sg, "total": total}


def extraer_semestre(lineas: list[str]) -> int | None:
    for i, ln in enumerate(lineas):
        n = normalizar(ln)
        m = re.search(r"semestre\s*[:\-]?\s*(\d+)", n)
        if m:
            return int(m.group(1))
        if "semestre" in n:
            prox = buscar_primer_numero(lineas, i + 1, alcance=4)
            if prox is not None and 1 <= prox <= 12:
                return prox
    return None


def extraer_prerrequisitos(lineas: list[str]) -> list[str]:
    prerreq: list[str] = []
    patron = re.compile(r"\b([A-Za-z]{2,10}\s?-?\d{3})\b")

    for i, ln in enumerate(lineas):
        n = normalizar(ln)
        if "prerrequis" in n or "pre requis" in n:
            ventana = [ln] + lineas[i + 1 : i + 4]
            joined = " ".join(ventana)
            for bruto in patron.findall(joined):
                norm = normalizar_codigo(bruto)
                if norm in CODIGO_NORMALIZADO:
                    prerreq.append(CODIGO_NORMALIZADO[norm])
    return deduplicar_preservando_orden(prerreq)


def extraer_bibliografia(lineas: list[str]) -> list[str]:
    start = None
    for i, ln in enumerate(lineas):
        n = normalizar(ln)
        if "bibliograf" in n or "referencias bibliograficas" in n:
            start = i + 1
            break
    if start is None:
        return []

    end = len(lineas)
    for i in range(start, len(lineas)):
        n = normalizar(lineas[i])
        if "detalle de las unidades" in n or n == "unidad tematica":
            end = i
            break

    referencias: list[str] = []
    for ln in lineas[start:end]:
        t = ln.strip()
        if len(t) >= 12 and re.search(r"\d{4}", t):
            referencias.append(t)
        elif re.match(r"^\d+[.)]\s+", t):
            referencias.append(t)
    return deduplicar_preservando_orden(referencias)


def _encontrar_tabla_detalle(doc: Document) -> Any | None:
    """Devuelve la tabla 'Detalle de las Unidades Temáticas', o None si no existe."""
    for tabla in doc.tables:
        for fila in tabla.rows[:3]:
            texto = " ".join(c.text for c in fila.cells)
            n = normalizar(texto)
            if "detalle" in n and "unidad" in n:
                return tabla
    return None


def _es_fila_total(celdas: list[str]) -> bool:
    """True si CUALQUIER celda de la fila contiene la palabra 'TOTAL'."""
    for c in celdas:
        nc = normalizar(c)
        if nc and re.search(r"\btotal\b", nc):
            return True
    return False


def _sumar_horas_fila(celdas: list[str], cols: list[int | None]) -> int:
    """Suma los valores numéricos de las columnas indicadas (ignora None)."""
    total = 0
    for col_idx in cols:
        if col_idx is not None and col_idx < len(celdas):
            try:
                total += int(celdas[col_idx] or "0")
            except ValueError:
                pass
    return total


def _encontrar_cols_horas(tabla: Any) -> tuple[int | None, int | None, int | None]:
    """Detecta qué columnas son T, P y AE/SG buscando esas etiquetas en las
    primeras 5 filas de la tabla.  Para tablas con celdas fusionadas, devuelve
    el PRIMER índice de columna que coincide con cada etiqueta.
    """
    col_t = col_p = col_ae = None
    for fila in tabla.rows[:5]:
        celdas = [c.text.strip() for c in fila.cells]
        for ci, val in enumerate(celdas):
            nv = normalizar(val)
            if nv == "t" and col_t is None:
                col_t = ci
            if nv == "p" and col_p is None:
                col_p = ci
            if nv in ("ae", "sg") and col_ae is None:
                col_ae = ci
        if col_t is not None and col_p is not None and col_ae is not None:
            break
    return col_t, col_p, col_ae


def extraer_unidades_tabla(
    doc: Document,
) -> tuple[list[dict[str, Any]], list[str], dict[str, int] | None]:
    """Extrae unidades temáticas de la tabla de unidades del Word.

    La estructura de columnas varía entre documentos:
      • 7 cols: col0=tipo | col1=RA | col2=Contenido | col3=Act | col4=T | col5=P | col6=AE
      • 9 cols: celdas fusionadas en T y P (2 columnas c/u)
      • 11 cols: celdas fusionadas en RA, Actividad, T y P (2 columnas c/u)

    Usa `_encontrar_cols_horas` para detectar los índices correctos de T, P, AE
    en lugar de asumir cols [4,5,6] (que falla en tablas de 9 y 11 columnas).

    Returns:
        (unidades, resultados_globales, horas_total_dict)
        horas_total_dict: {"T": t, "P": p, "SG": ae, "total": t+p+ae}
                          extraído de la fila TOTAL al final de la tabla,
                          o None si no se detectó ninguna fila TOTAL.
    """
    tabla = _encontrar_tabla_detalle(doc)
    if tabla is None:
        return [], [], None

    col_t, col_p, col_ae = _encontrar_cols_horas(tabla)
    cols_horas = [col_t, col_p, col_ae]

    unidades: list[dict[str, Any]] = []
    resultados_globales: list[str] = []
    unidad_actual: dict[str, Any] | None = None
    contenidos_vistos: set[str] = set()
    resultados_vistos: set[str] = set()
    actividades_vistas: set[str] = set()
    horas_total_dict: dict[str, int] | None = None

    for fila in tabla.rows:
        celdas = [c.text.strip() for c in fila.cells]
        if not any(celdas):
            continue

        n0 = normalizar(celdas[0])

        # ── Nueva unidad ──────────────────────────────────────────────────────
        if n0 in {"unidad tematica", "unidad temática"}:
            if unidad_actual is not None:
                unidades.append(unidad_actual)
            nombre = celdas[1] if len(celdas) > 1 else f"Unidad {len(unidades) + 1}"
            unidad_actual = {
                "nombre_unidad": nombre,
                "horas_asignadas": 0,
                "contenidos_especificos": [],
                "actividades_asociadas": [],
            }
            contenidos_vistos = set()
            resultados_vistos = set()
            actividades_vistas = set()
            continue

        # ── Fila TOTAL: extraer T/P/AE por separado y NO acumular ────────────
        if _es_fila_total(celdas):
            def _vi(idx: int | None) -> int:
                if idx is None or idx >= len(celdas):
                    return 0
                try:
                    return int(celdas[idx] or "0")
                except ValueError:
                    return 0
            t = _vi(col_t); p = _vi(col_p); ae = _vi(col_ae)
            if t + p + ae > 0:
                horas_total_dict = {"T": t, "P": p, "SG": ae, "total": t + p + ae}
            continue  # NO acumular a unidad_actual

        # ── Fila de cabecera de columnas o título de sección ─────────────────
        if unidad_actual is None:
            continue
        n1 = normalizar(celdas[1]) if len(celdas) > 1 else ""
        if "resultado" in n1 and n0 in {"", "ae", "t", "p"}:
            continue
        if "detalle" in n0 and "unidad" in n0:
            continue

        # ── Fila de datos ─────────────────────────────────────────────────────
        # Contenido — col 2
        if len(celdas) >= 3:
            for item in celdas[2].splitlines():
                item = item.strip()
                key = normalizar(item)
                if len(item) >= 8 and key not in contenidos_vistos:
                    contenidos_vistos.add(key)
                    unidad_actual["contenidos_especificos"].append(item)

        # Resultado de aprendizaje — col 1
        if len(celdas) >= 2:
            resultado = celdas[1]
            key = normalizar(resultado)
            if len(resultado) >= 8 and key not in resultados_vistos:
                resultados_vistos.add(key)
                resultados_globales.append(resultado)

        # Actividad — col 3
        if len(celdas) >= 4:
            actividad = celdas[3]
            key = normalizar(actividad)
            if len(actividad) >= 8 and key not in actividades_vistas:
                actividades_vistas.add(key)
                unidad_actual["actividades_asociadas"].append(actividad)

        # Horas de la unidad — usa los índices correctos de T, P, AE
        unidad_actual["horas_asignadas"] += _sumar_horas_fila(celdas, cols_horas)

    if unidad_actual is not None:
        unidades.append(unidad_actual)

    return unidades, list(dict.fromkeys(resultados_globales)), horas_total_dict


def procesar_programa(path_docx: Path) -> dict[str, Any]:
    doc = Document(path_docx)
    lineas = extraer_lineas_documento(doc)

    codigo   = extraer_codigo(path_docx, lineas)
    nombre   = extraer_nombre(path_docx, lineas, codigo)
    semestre = extraer_semestre(lineas)
    prerrequisitos = extraer_prerrequisitos(lineas)
    unidades, resultados, horas_total_dict = extraer_unidades_tabla(doc)
    bibliografia = extraer_bibliografia(lineas)

    # ── Fuente de horas: fila TOTAL de la tabla de unidades (preferida) ──────
    # La tabla resumen superior del Word tiene un bug estructural: la función
    # extraer_horas() encuentra el Total (mismo número) para T, P y AE porque
    # la tabla es horizontal y el primer número tras cada label es siempre el
    # Total.  La fila TOTAL al pie de la tabla de unidades es la fuente correcta.
    if horas_total_dict is not None:
        horas = horas_total_dict
        fuente_horas = "fila_TOTAL_tabla_unidades"
    else:
        horas = extraer_horas(lineas)  # fallback a tabla resumen superior
        fuente_horas = "tabla_resumen_superior"

    # ── Diagnóstico por asignatura ────────────────────────────────────────────
    etiqueta = codigo or path_docx.stem
    if unidades:
        suma_u = sum(u["horas_asignadas"] for u in unidades)
        partes = "+".join(str(u["horas_asignadas"]) for u in unidades)
        total_word = horas.get("total", 0)
        ok = "OK" if suma_u == total_word else f"DIFERENCIA (suma={suma_u} vs word={total_word})"
        print(f"  {etiqueta}: T={horas['T']} P={horas['P']} AE={horas['SG']} Total={total_word}h "
              f"[fuente: {fuente_horas}] [{ok}]")
    else:
        print(f"  {etiqueta}: sin unidades extraídas")

    faltantes: list[str] = []
    if not codigo:
        faltantes.append("codigo_asignatura")
    if not nombre:
        faltantes.append("nombre_asignatura")
    if horas.get("total", 0) == 0:
        faltantes.append("horas_totales")
    if semestre is None:
        faltantes.append("semestre")
    if not unidades:
        faltantes.append("unidades_tematicas")
    if not resultados:
        faltantes.append("resultados_aprendizaje")
    if not bibliografia:
        faltantes.append("bibliografia")

    return {
        "archivo_fuente": path_docx.name,
        "codigo_asignatura": codigo,
        "nombre_asignatura": nombre,
        "horas": horas,
        "fuente_horas": fuente_horas,
        "semestre": semestre,
        "prerrequisitos": prerrequisitos,
        "unidades_tematicas": unidades,
        "resultados_aprendizaje": resultados,
        "bibliografia": bibliografia,
        "campos_no_encontrados": faltantes,
    }


def _nc(s: str) -> str:
    return re.sub(r"[^a-z0-9]", "", s.lower())


def main() -> None:
    if not PROGRAMAS_DIR.exists():
        raise FileNotFoundError(f"No existe la carpeta de programas: {PROGRAMAS_DIR}")

    archivos = sorted(PROGRAMAS_DIR.glob("*.docx"))
    if not archivos:
        raise FileNotFoundError(f"No se encontraron archivos .docx en: {PROGRAMAS_DIR}")

    # Cargar Excel Tabla Resumen para comparación
    curriculum_path = BASE_DIR / "curriculum_data.json"
    ha_excel: dict[str, dict] = {}
    if curriculum_path.exists():
        curr = json.loads(curriculum_path.read_text(encoding="utf-8"))
        for cod_raw, v in curr.get("horas_asignaturas", {}).items():
            ha_excel[_nc(cod_raw)] = v

    print("=== EXTRACCION DE PROGRAMAS ===\n")
    programas: list[dict[str, Any]] = []
    for archivo in archivos:
        try:
            prog = procesar_programa(archivo)
            programas.append(prog)
        except Exception as exc:
            programas.append(
                {
                    "archivo_fuente": archivo.name,
                    "codigo_asignatura": extraer_codigo(archivo, []),
                    "error": f"Error procesando archivo: {exc}",
                    "campos_no_encontrados": ["documento_no_procesable"],
                }
            )

    salida = {
        "metadata": {
            "total_programas": len(programas),
            "directorio_origen": str(PROGRAMAS_DIR),
            "archivo_salida": str(OUTPUT_PATH),
        },
        "programas": programas,
    }

    OUTPUT_PATH.write_text(
        json.dumps(salida, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    print(f"\n{len(programas)} programas procesados -> {OUTPUT_PATH.name}")

    # ── Tabla de comparación Word vs Excel ────────────────────────────────────
    print("\n=== COMPARACIÓN WORD (fila TOTAL) vs EXCEL (Tabla Resumen) ===\n")
    enc = f"  {'Código':12s} │ {'T Word/Excel':>14} │ {'P Word/Excel':>14} │ {'AE Word/Excel':>15} │ {'Total':>5} │ Estado"
    sep = "─" * len(enc)
    print(enc)
    print(sep)

    ok_count = diff_count = err_count = 0
    for p in programas:
        cod = p.get("codigo_asignatura") or "?"
        if p.get("error"):
            print(f"  {cod:12s} │ ERROR: {p['error'][:60]}")
            err_count += 1
            continue
        h = p.get("horas", {})
        tw, pw, aw = h.get("T", 0), h.get("P", 0), h.get("SG", 0)
        tot_w = h.get("total", 0)

        xls = ha_excel.get(_nc(cod), {})
        diffs: list[str] = []
        if xls:
            tx  = int(xls.get("T", 0) or 0)
            px  = int(xls.get("P", 0) or 0)
            ax  = int(xls.get("SG", 0) or 0)
            matches = (tw == tx and pw == px and aw == ax)
            estado  = "✅" if matches else "⚠️"
            if matches:
                ok_count += 1
            else:
                diff_count += 1
                if tw != tx: diffs.append(f"T:{tw}→{tx}")
                if pw != px: diffs.append(f"P:{pw}→{px}")
                if aw != ax: diffs.append(f"AE:{aw}→{ax}")
        else:
            tx = px = ax = None
            matches = False
            estado  = "—"
            diff_count += 1
        nota = f"  ({', '.join(diffs)})" if diffs else ""

        def _c(w: int, x: int | None) -> str:
            if x is None: return f"{w:>4}/?"
            mark = " " if w == x else "*"
            return f"{w:>4}/{x:<4}{mark}"

        print(f"  {cod:12s} │ T:{_c(tw,tx)} │ P:{_c(pw,px)} │ AE:{_c(aw,ax)} │ Tot:{tot_w:>3} │ {estado}{nota}")

    print(sep)
    print(f"\n  {ok_count}/14 asignaturas con horas correctas ✅")
    print(f"  {diff_count}/14 asignaturas con diferencias  ⚠️")
    if err_count:
        print(f"  {err_count}/14 con error de procesamiento   ❌")

    errores = [p for p in programas if p.get("error")]
    if errores:
        print(f"\nERRORES ({len(errores)}):")
        for p in errores:
            print(f"  {p.get('codigo_asignatura') or p['archivo_fuente']}: {p['error']}")

    faltantes_resumen = [(p.get("codigo_asignatura") or p["archivo_fuente"],
                          p.get("campos_no_encontrados", []))
                         for p in programas if p.get("campos_no_encontrados")]
    if faltantes_resumen:
        print("\nCampos no extraidos:")
        for cod, falt in faltantes_resumen:
            print(f"  {cod}: {', '.join(falt)}")


if __name__ == "__main__":
    main()
