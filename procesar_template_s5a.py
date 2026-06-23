"""
Procesamiento del template IHO S-5A: extracción, validación y análisis.
Tareas 1-5 según especificación.
"""
from __future__ import annotations

import json
import re
import shutil
import sys
import tempfile
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document

BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "S-5A_Ed1.0.3-template (1).docx"
PROGRAMAS_PATH = BASE_DIR / "programas_estructurados.json"
OUTPUT_TEMPLATE = BASE_DIR / "template_s5a_estructurado.json"
OUTPUT_MAPA = BASE_DIR / "mapa_topico_asignatura.json"
OUTPUT_NIVELES = BASE_DIR / "niveles_iho_norma.json"
OUTPUT_PROFUNDIDAD = BASE_DIR / "profundidad_iho.json"

NIVEL_ORD = {"B": 1, "I": 2, "A": 3}

ASIG_NORMALIZE = {
    "PhyOce104": "PhyOce104",
}

ASIG_TYPO_FIX = {
    "RS202": "RS201",
    "SG303": "SG304",
    "WTC202": "WTC205",
}


def normalize_asig_code(raw: str) -> str:
    """Normaliza código de asignatura y corrige errores de tipeo del template."""
    code = ASIG_NORMALIZE.get(raw, raw.upper() if raw.isascii() else raw)
    return ASIG_TYPO_FIX.get(code, code)


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def save_json(path: Path, data: dict) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def abrir_documento() -> Document:
    try:
        return Document(str(TEMPLATE_PATH))
    except Exception:
        tmp = Path(tempfile.gettempdir()) / "S5A_template_copy.docx"
        shutil.copy2(str(TEMPLATE_PATH), str(tmp))
        return Document(str(tmp))


def non_empty(lines: list[str]) -> list[str]:
    return [l.strip() for l in lines if l.strip()]


def parse_topic_element(text: str) -> tuple[str, str, str]:
    """Extrae (código, nivel, descripción) de la columna Topic/Element."""
    full = text.strip()
    nivel_match = re.search(r'\(([BIA](?:\s*,\s*[BIA])*)\)\s*$', full)
    nivel = ""
    if nivel_match:
        niveles = [n.strip() for n in nivel_match.group(1).split(',')]
        nivel = max(niveles, key=lambda n: NIVEL_ORD.get(n, 0))
        full = full[:nivel_match.start()].strip()

    first_line = full.split('\n')[0].strip()
    code_match = re.match(r'([A-Z]\d+\.\d+[a-z]?)\s*(.*)', first_line)
    if code_match:
        return code_match.group(1), nivel, code_match.group(2).strip()
    return first_line, nivel, ""


def parse_module_entry(entry: str) -> dict:
    """Parsea una entrada individual de Module & Content."""
    entry = entry.strip()

    # Pattern with T.U. and optional contents
    m = re.match(
        r'([A-Za-z]{2,6})\s*(\d{3})\s+'
        r'T\.?U\.?\s*(\d+\.\d+)\s*(\([^)]*\))?\s*$',
        entry
    )
    if m:
        asig = f"{m.group(1)}{m.group(2)}"
        return {
            "asignatura": normalize_asig_code(asig),
            "tu": m.group(3),
            "contenidos": m.group(4) or "",
            "detalle": "",
            "raw": entry,
        }

    # General pattern: CODE NNN followed by anything
    m2 = re.match(r'([A-Za-z]{2,6})\s*(\d{3})\s+(.*)', entry)
    if m2:
        asig = f"{m2.group(1)}{m2.group(2)}"
        rest = m2.group(3).strip()
        tu_m = re.search(r'T\.?U\.?\s*(\d+\.\d+)', rest)
        tu = tu_m.group(1) if tu_m else ""
        cont_m = re.search(r'(\([^)]*\))\s*$', rest)
        contenidos = cont_m.group(1) if cont_m else ""
        detalle_text = rest
        if tu_m:
            detalle_text = detalle_text[:tu_m.start()] + detalle_text[tu_m.end():]
        if cont_m:
            detalle_text = detalle_text[:cont_m.start()]
        detalle_text = detalle_text.strip().strip(',').strip()
        return {
            "asignatura": normalize_asig_code(asig),
            "tu": tu,
            "contenidos": contenidos,
            "detalle": detalle_text,
            "raw": entry,
        }

    # Minimal: just CODE NNN
    m3 = re.match(r'([A-Za-z]{2,6})\s*(\d{3})', entry)
    if m3:
        asig = f"{m3.group(1)}{m3.group(2)}"
        return {
            "asignatura": normalize_asig_code(asig),
            "tu": "",
            "contenidos": "",
            "detalle": entry[m3.end():].strip(),
            "raw": entry,
        }

    return None


def es_fila_datos(cells) -> bool:
    c0 = cells[0].text.strip()
    c3 = cells[3].text.strip()
    if not c0:
        return False
    if c3 in ('TOTAL', 'OVERALL TOTAL', 'TOTAL PROGRAM HOURS', ''):
        return False
    if cells[0].text.strip() == cells[1].text.strip() == cells[2].text.strip():
        return False
    if 'Topic/Element' in c0:
        return False
    return True


def extraer_tabla(table) -> list[dict]:
    """Extrae todas las filas de datos de una tabla del template."""
    filas = []
    topico_padre = ""
    seccion = ""

    for row in table.rows:
        cells = row.cells
        c0 = cells[0].text.strip()

        if cells[0].text.strip() == cells[1].text.strip() == cells[2].text.strip():
            sec_match = re.match(r'([A-Z]\d+):', c0)
            if sec_match:
                seccion = sec_match.group(1)
            sub_match = re.match(r'([A-Z]\d+\.\d+)\s', c0)
            if sub_match:
                topico_padre = sub_match.group(1)
            continue

        if not es_fila_datos(cells):
            continue

        code, nivel, desc_topic = parse_topic_element(cells[0].text)
        content = cells[1].text.strip()
        learning_outcomes = cells[2].text.strip()

        mod_entries = non_empty(cells[3].text.split('\n'))
        t_entries = non_empty(cells[4].text.split('\n'))
        p_entries = non_empty(cells[5].text.split('\n'))
        sg_entries = non_empty(cells[6].text.split('\n'))

        n = len(mod_entries)
        while len(t_entries) < n:
            t_entries.append("0")
        while len(p_entries) < n:
            p_entries.append("0")
        while len(sg_entries) < n:
            sg_entries.append("0")

        for i in range(n):
            mod = parse_module_entry(mod_entries[i])
            if not mod:
                continue

            def to_int(s):
                s = s.strip()
                try:
                    return int(s)
                except ValueError:
                    return 0

            filas.append({
                "codigo_topico_iho": code,
                "topico_padre": topico_padre,
                "seccion": seccion,
                "nivel_iho": nivel,
                "descripcion_topico": desc_topic,
                "contenido_norma": content,
                "learning_outcome": learning_outcomes,
                "asignatura": mod["asignatura"],
                "unidad_tematica": mod["tu"],
                "contenidos_ref": mod["contenidos"],
                "detalle": mod["detalle"],
                "modulo_raw": mod["raw"],
                "horas_T": to_int(t_entries[i]),
                "horas_P": to_int(p_entries[i]),
                "horas_SG": to_int(sg_entries[i]),
            })

    return filas


# ═══════════════════════════════════════════════════════════════════
# TAREA 1
# ═══════════════════════════════════════════════════════════════════
def tarea1_leer_estructurar(doc: Document) -> list[dict]:
    print("=" * 70)
    print("TAREA 1: LEER Y ESTRUCTURAR EL TEMPLATE S-5A")
    print("=" * 70)

    print(f"\nNúmero de tablas: {len(doc.tables)}")
    for i, t in enumerate(doc.tables):
        ncols = len(t.rows[0].cells) if t.rows else 0
        print(f"  Tabla {i}: {len(t.rows)} filas x {ncols} columnas")

    filas_t0 = extraer_tabla(doc.tables[0])
    filas_t1 = extraer_tabla(doc.tables[1])
    todas = filas_t0 + filas_t1

    print(f"\nFilas extraídas tabla 0 (Foundation): {len(filas_t0)}")
    print(f"Filas extraídas tabla 1 (Hydrographic): {len(filas_t1)}")
    print(f"Total filas de datos: {len(todas)}")

    asignaturas = sorted(set(f["asignatura"] for f in todas))
    topicos = sorted(set(f["codigo_topico_iho"] for f in todas))
    print(f"Asignaturas únicas: {len(asignaturas)} → {asignaturas}")
    print(f"Tópicos IHO únicos: {len(topicos)}")

    total_T = sum(f["horas_T"] for f in todas)
    total_P = sum(f["horas_P"] for f in todas)
    total_SG = sum(f["horas_SG"] for f in todas)
    print(f"Horas totales template: T={total_T} P={total_P} SG={total_SG} Total={total_T+total_P+total_SG}")

    salida = {
        "metadata": {
            "fuente": TEMPLATE_PATH.name,
            "total_filas": len(todas),
            "total_asignaturas": len(asignaturas),
            "total_topicos_iho": len(topicos),
            "horas_globales": {"T": total_T, "P": total_P, "SG": total_SG, "total": total_T + total_P + total_SG},
            "asignaturas": asignaturas,
        },
        "filas": todas,
    }
    save_json(OUTPUT_TEMPLATE, salida)
    print(f"\nExportado a: {OUTPUT_TEMPLATE.name}")

    print(f"\n--- Primeras 10 filas (verificación) ---")
    for i, f in enumerate(todas[:10]):
        total = f["horas_T"] + f["horas_P"] + f["horas_SG"]
        print(f"  [{i}] {f['codigo_topico_iho']} ({f['nivel_iho']}) → "
              f"{f['asignatura']} TU {f['unidad_tematica']} {f['contenidos_ref']} "
              f"| T={f['horas_T']} P={f['horas_P']} SG={f['horas_SG']} ={total}")

    return todas


# ═══════════════════════════════════════════════════════════════════
# TAREA 2
# ═══════════════════════════════════════════════════════════════════
def tarea2_validacion_horas(filas: list[dict]) -> list[dict]:
    print(f"\n{'=' * 70}")
    print("TAREA 2: VALIDACIÓN DE HORAS (Template vs Programas Word)")
    print("=" * 70)

    programas = load_json(PROGRAMAS_PATH)
    horas_word = {}
    for p in programas.get("programas", []):
        cod = p.get("codigo_asignatura", "")
        if cod and not p.get("error"):
            horas_word[cod] = p["horas"]

    horas_template: dict[str, dict[str, int]] = {}
    for f in filas:
        asig = f["asignatura"]
        if not asig:
            continue
        if asig not in horas_template:
            horas_template[asig] = {"T": 0, "P": 0, "SG": 0}
        horas_template[asig]["T"] += f["horas_T"]
        horas_template[asig]["P"] += f["horas_P"]
        horas_template[asig]["SG"] += f["horas_SG"]

    all_asigs = sorted(set(list(horas_template.keys()) + list(horas_word.keys())))

    print(f"\n{'Asignatura':<12} {'T_tmpl':>6} {'T_Word':>6} {'DifT':>5} "
          f"{'P_tmpl':>6} {'P_Word':>6} {'DifP':>5} "
          f"{'SG_tmpl':>7} {'SG_Word':>7} {'DifSG':>5} "
          f"{'Tot_t':>5} {'Tot_W':>5} {'Estado'}")
    print("-" * 110)

    resultados = []
    for asig in all_asigs:
        ht = horas_template.get(asig, {"T": 0, "P": 0, "SG": 0})
        hw = horas_word.get(asig, {"T": 0, "P": 0, "SG": 0, "total": 0})

        diff_t = ht["T"] - hw.get("T", 0)
        diff_p = ht["P"] - hw.get("P", 0)
        diff_sg = ht["SG"] - hw.get("SG", 0)
        total_t = ht["T"] + ht["P"] + ht["SG"]
        total_w = hw.get("total", hw.get("T", 0) + hw.get("P", 0) + hw.get("SG", 0))
        diff_total = abs(total_t - total_w)

        if diff_t == 0 and diff_p == 0 and diff_sg == 0:
            estado = "✅"
        elif diff_total <= 3:
            estado = "⚠️"
        else:
            estado = "❌"

        print(f"{asig:<12} {ht['T']:>6} {hw.get('T',0):>6} {diff_t:>+5} "
              f"{ht['P']:>6} {hw.get('P',0):>6} {diff_p:>+5} "
              f"{ht['SG']:>7} {hw.get('SG',0):>7} {diff_sg:>+5} "
              f"{total_t:>5} {total_w:>5} {estado}")

        resultados.append({
            "asignatura": asig,
            "T_template": ht["T"], "T_Word": hw.get("T", 0), "diff_T": diff_t,
            "P_template": ht["P"], "P_Word": hw.get("P", 0), "diff_P": diff_p,
            "SG_template": ht["SG"], "SG_Word": hw.get("SG", 0), "diff_SG": diff_sg,
            "total_template": total_t, "total_Word": total_w,
            "estado": estado,
            "en_template": asig in horas_template,
            "en_word": asig in horas_word,
        })

    n_ok = sum(1 for r in resultados if r["estado"] == "✅")
    n_warn = sum(1 for r in resultados if r["estado"] == "⚠️")
    n_err = sum(1 for r in resultados if r["estado"] == "❌")
    solo_template = [r["asignatura"] for r in resultados if not r["en_word"]]
    solo_word = [r["asignatura"] for r in resultados if not r["en_template"]]

    print(f"\nResumen: ✅ {n_ok} coinciden | ⚠️ {n_warn} leve | ❌ {n_err} revisar")
    if solo_template:
        print(f"Solo en template (no en Word): {solo_template}")
    if solo_word:
        print(f"Solo en Word (no en template): {solo_word}")

    return resultados


# ═══════════════════════════════════════════════════════════════════
# TAREA 3
# ═══════════════════════════════════════════════════════════════════
def tarea3_niveles_reales(filas: list[dict]) -> dict:
    print(f"\n{'=' * 70}")
    print("TAREA 3: NIVELES IHO REALES DESDE EL TEMPLATE")
    print("=" * 70)

    elementos = {}
    for f in filas:
        code = f["codigo_topico_iho"]
        nivel = f["nivel_iho"]
        if not code or not nivel:
            continue
        if code not in elementos:
            elementos[code] = nivel
        elif NIVEL_ORD.get(nivel, 0) > NIVEL_ORD.get(elementos[code], 0):
            elementos[code] = nivel

    topicos = {}
    for code, nivel in elementos.items():
        padre = re.match(r'([A-Z]\d+\.\d+)', code)
        if padre:
            p = padre.group(1)
            if p not in topicos:
                topicos[p] = nivel
            elif NIVEL_ORD.get(nivel, 0) > NIVEL_ORD.get(topicos[p], 0):
                topicos[p] = nivel

    niveles_data = {
        "metadata": {
            "fuente": TEMPLATE_PATH.name,
            "nota": "Niveles de competencia IHO S-5A extraídos del template oficial: B=Basic, I=Intermediate, A=Advanced.",
            "total_elementos": len(elementos),
            "total_topicos_padre": len(topicos),
        },
        "elementos": dict(sorted(elementos.items())),
        "topicos": dict(sorted(topicos.items())),
    }

    save_json(OUTPUT_NIVELES, niveles_data)

    print(f"\nElementos con nivel B/I/A: {len(elementos)}")
    print(f"Tópicos padre: {len(topicos)}")

    dist_elem = {"B": 0, "I": 0, "A": 0}
    for n in elementos.values():
        dist_elem[n] = dist_elem.get(n, 0) + 1
    print(f"Distribución elementos: B={dist_elem['B']} I={dist_elem['I']} A={dist_elem['A']}")

    dist_top = {"B": 0, "I": 0, "A": 0}
    for n in topicos.values():
        dist_top[n] = dist_top.get(n, 0) + 1
    print(f"Distribución tópicos:   B={dist_top['B']} I={dist_top['I']} A={dist_top['A']}")

    print(f"\nActualizado: {OUTPUT_NIVELES.name}")
    return niveles_data


# ═══════════════════════════════════════════════════════════════════
# TAREA 4
# ═══════════════════════════════════════════════════════════════════
def tarea4_mapa_topico_asignatura(filas: list[dict]) -> list[dict]:
    print(f"\n{'=' * 70}")
    print("TAREA 4: MAPA COMPLETO TÓPICO → ASIGNATURA")
    print("=" * 70)

    mapa = []
    for f in filas:
        if not f["asignatura"]:
            continue
        total = f["horas_T"] + f["horas_P"] + f["horas_SG"]
        mapa.append({
            "topico_iho": f["codigo_topico_iho"],
            "nivel_BIA": f["nivel_iho"],
            "asignatura_shoa": f["asignatura"],
            "unidad_tematica": f["unidad_tematica"],
            "contenidos_ref": f["contenidos_ref"],
            "horas_T": f["horas_T"],
            "horas_P": f["horas_P"],
            "horas_SG": f["horas_SG"],
            "total_horas": total,
        })

    save_json(OUTPUT_MAPA, {
        "metadata": {
            "fuente": TEMPLATE_PATH.name,
            "total_registros": len(mapa),
            "descripcion": "Mapa completo tópico IHO → asignatura SHOA con horas",
        },
        "mapa": mapa,
    })

    # Tabla por asignatura
    por_asig: dict[str, dict] = {}
    for m in mapa:
        asig = m["asignatura_shoa"]
        if asig not in por_asig:
            por_asig[asig] = {"topicos": set(), "T": 0, "P": 0, "SG": 0}
        por_asig[asig]["topicos"].add(m["topico_iho"])
        por_asig[asig]["T"] += m["horas_T"]
        por_asig[asig]["P"] += m["horas_P"]
        por_asig[asig]["SG"] += m["horas_SG"]

    print(f"\nTotal registros en mapa: {len(mapa)}")
    print(f"\n{'Asignatura':<12} {'#Tópicos':>8} {'T':>5} {'P':>5} {'SG':>5} {'Total':>6}")
    print("-" * 50)
    for asig in sorted(por_asig):
        d = por_asig[asig]
        total = d["T"] + d["P"] + d["SG"]
        print(f"{asig:<12} {len(d['topicos']):>8} {d['T']:>5} {d['P']:>5} {d['SG']:>5} {total:>6}")

    # Tabla por tópico
    por_topico: dict[str, dict] = {}
    for m in mapa:
        t = m["topico_iho"]
        if t not in por_topico:
            por_topico[t] = {"nivel": m["nivel_BIA"], "asignaturas": set(), "T": 0, "P": 0, "SG": 0}
        por_topico[t]["asignaturas"].add(m["asignatura_shoa"])
        por_topico[t]["T"] += m["horas_T"]
        por_topico[t]["P"] += m["horas_P"]
        por_topico[t]["SG"] += m["horas_SG"]

    print(f"\n{'Tópico':<12} {'Nivel':>5} {'Asignaturas':<40} {'T':>4} {'P':>4} {'SG':>4} {'Tot':>5}")
    print("-" * 82)
    for top in sorted(por_topico):
        d = por_topico[top]
        asigs = ", ".join(sorted(d["asignaturas"]))
        total = d["T"] + d["P"] + d["SG"]
        print(f"{top:<12} {d['nivel']:>5} {asigs:<40} {d['T']:>4} {d['P']:>4} {d['SG']:>4} {total:>5}")

    print(f"\nExportado a: {OUTPUT_MAPA.name}")
    return mapa


# ═══════════════════════════════════════════════════════════════════
# TAREA 5
# ═══════════════════════════════════════════════════════════════════
def tarea5_actualizar_analisis(filas: list[dict], niveles: dict) -> None:
    print(f"\n{'=' * 70}")
    print("TAREA 5: ACTUALIZAR ANÁLISIS Y RESUMEN FINAL")
    print("=" * 70)

    topicos_con_nivel = set()
    for f in filas:
        if f["nivel_iho"]:
            topicos_con_nivel.add(f["codigo_topico_iho"])

    asignaturas_cubiertas = set(f["asignatura"] for f in filas if f["asignatura"])

    topicos_sin_asig = set()
    topicos_con_asig = set()
    for f in filas:
        if f["asignatura"]:
            topicos_con_asig.add(f["codigo_topico_iho"])
        elif f["codigo_topico_iho"]:
            topicos_sin_asig.add(f["codigo_topico_iho"])
    topicos_sin_asig -= topicos_con_asig

    programas = load_json(PROGRAMAS_PATH)
    asigs_word = set()
    horas_word = {}
    for p in programas.get("programas", []):
        cod = p.get("codigo_asignatura", "")
        if cod and not p.get("error"):
            asigs_word.add(cod)
            horas_word[cod] = p["horas"]

    asigs_no_en_template = asigs_word - asignaturas_cubiertas

    horas_template: dict[str, dict[str, int]] = {}
    for f in filas:
        asig = f["asignatura"]
        if not asig:
            continue
        if asig not in horas_template:
            horas_template[asig] = {"T": 0, "P": 0, "SG": 0}
        horas_template[asig]["T"] += f["horas_T"]
        horas_template[asig]["P"] += f["horas_P"]
        horas_template[asig]["SG"] += f["horas_SG"]

    discrepancias_graves = 0
    for asig in horas_template:
        if asig in horas_word:
            ht = horas_template[asig]
            hw = horas_word[asig]
            total_t = ht["T"] + ht["P"] + ht["SG"]
            total_w = hw.get("total", 0)
            if abs(total_t - total_w) > 3:
                discrepancias_graves += 1

    profundidad_filas = []
    for f in filas:
        if not f["asignatura"]:
            continue
        profundidad_filas.append({
            "codigo_topico_iho": f["codigo_topico_iho"],
            "topico_padre": f["topico_padre"],
            "nivel_iho_requerido": f["nivel_iho"],
            "asignatura": f["asignatura"],
            "unidad_tematica": f["unidad_tematica"],
            "horas_T": f["horas_T"],
            "horas_P": f["horas_P"],
            "horas_SG": f["horas_SG"],
        })

    save_json(OUTPUT_PROFUNDIDAD, {
        "metadata": {
            "fuente": TEMPLATE_PATH.name,
            "total_registros": len(profundidad_filas),
            "niveles_fuente": "template S-5A oficial (niveles reales B/I/A)",
        },
        "registros": profundidad_filas,
    })

    print(f"\n{'=' * 50}")
    print("RESUMEN FINAL")
    print("=" * 50)
    print(f"✅ {len(topicos_con_nivel)} tópicos con nivel B/I/A extraídos")
    print(f"📚 {len(asignaturas_cubiertas)} asignaturas cubiertas en el template")
    if topicos_sin_asig:
        print(f"⚠️  {len(topicos_sin_asig)} tópicos sin asignatura asignada: {sorted(topicos_sin_asig)}")
    else:
        print(f"⚠️  0 tópicos sin asignatura asignada")
    if asigs_no_en_template:
        print(f"📋 {len(asigs_no_en_template)} asignaturas Word no en template: {sorted(asigs_no_en_template)}")
    print(f"❌ {discrepancias_graves} discrepancias graves con Word TOTAL (>3h)")

    print(f"\nArchivos actualizados:")
    print(f"  → {OUTPUT_NIVELES.name}")
    print(f"  → {OUTPUT_PROFUNDIDAD.name}")
    print(f"  → {OUTPUT_TEMPLATE.name}")
    print(f"  → {OUTPUT_MAPA.name}")


def main():
    print("Procesamiento del template IHO S-5A")
    print(f"Archivo: {TEMPLATE_PATH.name}\n")

    doc = abrir_documento()

    filas = tarea1_leer_estructurar(doc)
    tarea2_validacion_horas(filas)
    niveles = tarea3_niveles_reales(filas)
    tarea4_mapa_topico_asignatura(filas)
    tarea5_actualizar_analisis(filas, niveles)


if __name__ == "__main__":
    main()
