"""Dashboard de Diseño Curricular SHOA — Fase 2."""
from __future__ import annotations

import io
import json
from collections import defaultdict
from datetime import date
from pathlib import Path
from typing import Any

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st

# ── Rutas ────────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent
ASSET_DIR = BASE_DIR / "assets"
LOGO_PATH = ASSET_DIR / "logo_shoa.png"

PATHS: dict[str, Path] = {
    "programas":   BASE_DIR / "programas_estructurados.json",
    "progresion":  BASE_DIR / "progresion_analisis.json",
    "reiteracion": BASE_DIR / "reiteracion_matriz.json",
    "profundidad": BASE_DIR / "profundidad_iho.json",
    "propuesta":   BASE_DIR / "propuesta_curricular.json",
    "curriculum":  BASE_DIR / "curriculum_data.json",
}

# ── Constantes ────────────────────────────────────────────────────────────────
NIVEL_COLOR: dict[int, str] = {
    1: "#1565C0",
    2: "#2E7D32",
    3: "#E65100",
    4: "#6A1B9A",
    5: "#B71C1C",
}
NIVEL_LABEL: dict[int, str] = {
    1: "N1 – Fundamentos (F1-F3)",
    2: "N2 – Instrumentación (H1-H3)",
    3: "N3 – Operaciones (H4-H5)",
    4: "N4 – Procesamiento (H6-H7)",
    5: "N5 – Gestión (H8)",
}
ACCION_COLOR: dict[str, str] = {
    "MANTENER":                    "#4CAF50",
    "REDUCIR":                     "#FFC107",
    "REDUCIR (Práctica de Campo)": "#8BC34A",
    "FUSIONAR":                    "#FF9800",
    "REORGANIZAR":                 "#2196F3",
    "ELIMINAR CONTENIDOS":         "#F44336",
}
CLASIF_COLOR: dict[str, str] = {
    "SOBREDIMENSIONADO": "#E53935",
    "ALINEADO":          "#43A047",
    "SUBDIMENSIONADO":   "#1E88E5",
}
# Escala IHO S-5A: B=Básico(1) < I=Intermedio(2) < A=Avanzado(3)
NIVEL_MAP: dict[str, int] = {"B": 1, "I": 2, "A": 3}
NIVEL_LABEL_BIA: dict[int, str] = {1: "B – Básico", 2: "I – Intermedio", 3: "A – Avanzado"}

# ── Carga de datos ────────────────────────────────────────────────────────────
def _leer(path_str: str) -> dict[str, Any] | None:
    p = Path(path_str)
    if not p.exists():
        return None
    return json.loads(p.read_text(encoding="utf-8"))


def cargar(key: str, mostrar_aviso: bool = True) -> dict[str, Any] | None:
    data = _leer(str(PATHS[key]))
    if data is None and mostrar_aviso:
        st.warning(
            f"⚠️ **{PATHS[key].name}** no encontrado. "
            "Ejecuta primero el script correspondiente."
        )
    return data


_DATOS_VERSION = "2026-06-16"


def _aviso_script(nombre: str) -> None:
    st.warning(
        f"⚠️ **{nombre.replace('.py','').replace('_',' ').title()}** "
        "no encontrado en los datos precargados. "
        "Contacta al administrador del dashboard."
    )


# ── Grafo interactivo ─────────────────────────────────────────────────────────
def _posiciones_semestre(
    nodos: list[str],
    nodo_info: dict[str, Any],
    override: dict[str, int] | None = None,
) -> dict[str, tuple[float, float]]:
    grupos: dict[int, list[str]] = defaultdict(list)
    for n in nodos:
        sem = override.get(n) if override else None
        if sem is None:
            sem = int(nodo_info.get(n, {}).get("semestre_actual", 1))
        grupos[sem].append(n)
    pos: dict[str, tuple[float, float]] = {}
    for sem, lista in grupos.items():
        for i, n in enumerate(lista):
            y = (i - (len(lista) - 1) / 2.0) * 2.2
            pos[n] = (float(sem) * 3.5, y)
    return pos


def fig_grafo(
    asignaturas: list[dict],
    dependencias: list[dict],
    aristas_problema: set[tuple[str, str]],
    override: dict[str, int] | None = None,
    titulo: str = "Grafo de Progresión",
) -> go.Figure:
    nodo_info = {a["codigo"]: a for a in asignaturas}
    nodos = list(nodo_info)
    pos = _posiciones_semestre(nodos, nodo_info, override)

    traces: list[Any] = []

    # Aristas
    for dep in dependencias:
        u, v = dep["desde"], dep["hacia"]
        if u not in pos or v not in pos:
            continue
        x0, y0 = pos[u]
        x1, y1 = pos[v]
        es_problema = (u, v) in aristas_problema
        traces.append(go.Scatter(
            x=[x0, x1, None], y=[y0, y1, None],
            mode="lines",
            line=dict(width=3 if es_problema else 1.5,
                      color="#E53935" if es_problema else "#BDBDBD"),
            hoverinfo="none", showlegend=False,
        ))

    # Nodos agrupados por nivel
    for nivel, color in NIVEL_COLOR.items():
        grupo = [n for n in nodos if n in pos
                 and nodo_info.get(n, {}).get("nivel_iho_predominante") == nivel]
        if not grupo:
            continue
        hovers = []
        for n in grupo:
            info = nodo_info[n]
            sem = override.get(n, info.get("semestre_actual", "?")) if override else info.get("semestre_actual", "?")
            hovers.append(
                f"<b>{n}</b><br>"
                f"{info.get('nombre', '')}<br>"
                f"Semestre: {sem} | Nivel IHO: {nivel}<br>"
                f"Prereqs: {', '.join(info.get('prerrequisitos', [])) or 'Ninguno'}"
            )
        traces.append(go.Scatter(
            x=[pos[n][0] for n in grupo],
            y=[pos[n][1] for n in grupo],
            mode="markers+text",
            marker=dict(size=42, color=color,
                        line=dict(width=2, color="white"), opacity=0.92),
            text=grupo,
            textfont=dict(size=8, color="white"),
            textposition="middle center",
            name=NIVEL_LABEL[nivel],
            hovertext=hovers, hoverinfo="text",
        ))

    # Etiquetas de semestre
    sems_presentes: set[int] = set()
    for n in pos:
        sem = override.get(n) if override else None
        if sem is None:
            sem = int(nodo_info.get(n, {}).get("semestre_actual", 1))
        sems_presentes.add(sem)

    annotations = [
        dict(x=float(s) * 3.5, y=4.8, text=f"<b>Semestre {s}</b>",
             showarrow=False, font=dict(size=13, color="#424242"), xanchor="center")
        for s in sorted(sems_presentes)
    ]

    fig = go.Figure(data=traces)
    fig.update_layout(
        title=dict(text=titulo, font=dict(size=15)),
        showlegend=True,
        hovermode="closest",
        xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        height=520,
        margin=dict(l=10, r=10, t=55, b=10),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        annotations=annotations,
        plot_bgcolor="rgba(245,245,245,0.5)",
    )
    return fig


# ── Tab 1 — Progresión ────────────────────────────────────────────────────────
def tab_progresion() -> None:
    st.header("🗺️ Mapa de Progresión")

    prog = cargar("progresion")
    if prog is None:
        _aviso_script("analisis_progresion.py")
        return

    asignaturas = prog.get("asignaturas", [])
    dependencias = prog.get("dependencias", [])
    problemas = prog.get("problemas_progresion", [])
    aristas_problema = {
        (a["desde"], a["hacia"]) for a in prog.get("aristas_con_problema", [])
    }

    prop_data = cargar("propuesta", mostrar_aviso=False)
    sem_propuesto: dict[str, int] = {}
    if prop_data:
        for p in prop_data.get("propuesta_asignaturas", []):
            sem_propuesto[p["asignatura"]] = p["semestre_propuesto"]

    # KPIs
    c1, c2, c3 = st.columns(3)
    c1.metric("Asignaturas analizadas", len(asignaturas))
    c2.metric("Dependencias mapeadas", len(dependencias))
    c3.metric("Problemas detectados", len(problemas))

    # Leyenda de colores
    cols_ley = st.columns(5)
    for i, (nivel, color) in enumerate(NIVEL_COLOR.items()):
        cols_ley[i].markdown(
            f"<div style='background:{color};color:white;padding:3px 6px;"
            f"border-radius:4px;text-align:center;font-size:11px'>"
            f"{NIVEL_LABEL[nivel]}</div>",
            unsafe_allow_html=True,
        )
    st.markdown("&nbsp;", unsafe_allow_html=True)

    # Grafos: actual vs propuesto
    if sem_propuesto:
        col_a, col_b = st.columns(2)
        with col_a:
            st.plotly_chart(
                fig_grafo(asignaturas, dependencias, aristas_problema,
                          titulo="Estructura Actual"),
                width="stretch",
            )
        with col_b:
            st.plotly_chart(
                fig_grafo(asignaturas, dependencias, aristas_problema,
                          override=sem_propuesto,
                          titulo="Estructura Propuesta (3 semestres)"),
                width="stretch",
            )
    else:
        st.plotly_chart(
            fig_grafo(asignaturas, dependencias, aristas_problema,
                      titulo="Estructura Actual"),
            width="stretch",
        )
        st.info("Ejecuta `propuesta_curricular.py` para ver la vista propuesta en paralelo.")

    st.caption("🔴 Aristas rojas = problemas de progresión detectados")

    # Lista de problemas
    st.markdown("---")
    st.subheader("Problemas de progresión detectados")
    if not problemas:
        st.success("✅ No se detectaron problemas de progresión.")
        return
    for i, pb in enumerate(problemas, 1):
        tipo = pb.get("tipo", "").replace("_", " ").title()
        with st.expander(f"🔴 Problema {i} — {tipo}", expanded=i <= 3):
            st.write(pb.get("descripcion", ""))
            ca, cb = st.columns(2)
            ca.write(f"**Asignatura:** `{pb.get('asignatura', '—')}`")
            cb.write(f"**Semestre:** {pb.get('semestre_asignatura', '—')}")
            if "asignatura_base" in pb:
                ca.write(f"**Asignatura base:** `{pb.get('asignatura_base', '—')}`")
                cb.write(f"**Semestre base:** {pb.get('semestre_base', '—')}")

    sin_prereq = prog.get("asignaturas_sin_prerrequisitos_claros", [])
    if sin_prereq:
        st.subheader("Asignaturas avanzadas sin prerrequisitos explícitos")
        df_sp = pd.DataFrame([
            {"Código": a["codigo"], "Nombre": a.get("nombre", ""), "Nivel IHO": a.get("nivel_iho_predominante")}
            for a in sin_prereq
        ])
        st.dataframe(df_sp, width="stretch")


# ── Tab 2 — Reiteración ───────────────────────────────────────────────────────
def tab_reiteracion() -> None:
    st.header("🔁 Matriz de Reiteración de Contenidos")

    reit = cargar("reiteracion")
    if reit is None:
        _aviso_script("matriz_reiteracion.py")
        return

    resumen = reit.get("resumen", {})
    reiteraciones = reit.get("reiteraciones", [])
    hm = reit.get("heatmap", {})

    # KPIs
    c1, c2, c3 = st.columns(3)
    c1.metric("Total reiteraciones", resumen.get("total_reiteraciones_detectadas", 0))
    c2.metric("Innecesarias", resumen.get("total_innecesarias", 0))
    c3.metric("Horas recuperables est.", f"{resumen.get('horas_recuperables_estimadas', 0):.1f} h")

    # Heatmap
    asigs_hm = hm.get("asignaturas", [])
    matriz = hm.get("matriz_total", [])
    if asigs_hm and matriz:
        st.subheader("Mapa de calor — contenidos compartidos entre asignaturas")
        fig_hm = px.imshow(
            matriz,
            x=asigs_hm, y=asigs_hm,
            color_continuous_scale="YlOrRd",
            labels=dict(color="N° contenidos"),
            aspect="auto",
        )
        fig_hm.update_layout(height=420)
        st.plotly_chart(fig_hm, width="stretch")

    # Filtro por tipo
    st.subheader("Tabla detallada de reiteraciones")
    tipos = sorted({r["tipo_reiteracion"] for r in reiteraciones})
    tipo_sel = st.multiselect("Filtrar por tipo", options=tipos, default=tipos)
    filtrado = [r for r in reiteraciones if r["tipo_reiteracion"] in tipo_sel]

    if filtrado:
        tipo_bg = {"INNECESARIA": "#ffcccc", "REVISAR": "#fff3cc", "JUSTIFICADA": "#d4edda"}

        df = pd.DataFrame([
            {
                "Contenido (Asig. 1)": r["contenido"][:70],
                "Asig. 1": r["asignatura_1"],
                "Sem. 1": r["semestre_1"],
                "N.IHO 1": r["nivel_iho_1"],
                "Contenido (Asig. 2)": r["contenido_relacionado"][:70],
                "Asig. 2": r["asignatura_2"],
                "Sem. 2": r["semestre_2"],
                "N.IHO 2": r["nivel_iho_2"],
                "Tipo": r["tipo_reiteracion"],
                "Similitud": r["similitud"],
                "Horas dup.": r["horas_duplicadas_estimadas"],
            }
            for r in filtrado
        ])
        st.dataframe(
            df.style.map(
                lambda v: f"background-color:{tipo_bg.get(v, 'white')}",
                subset=["Tipo"],
            ),
            width="stretch",
            height=420,
        )
    else:
        st.info("Sin reiteraciones con los filtros seleccionados.")

    # Gráfico de distribución por tipo
    if reiteraciones:
        conteos_tipo = pd.Series([r["tipo_reiteracion"] for r in reiteraciones]).value_counts().reset_index()
        conteos_tipo.columns = ["Tipo", "Cantidad"]
        fig_tipo = px.bar(
            conteos_tipo, x="Tipo", y="Cantidad",
            color="Tipo",
            color_discrete_map={"INNECESARIA": "#E53935", "REVISAR": "#FB8C00", "JUSTIFICADA": "#43A047"},
            title="Distribución por tipo de reiteración",
        )
        st.plotly_chart(fig_tipo, width="stretch")


# ── Tab 3 — Profundidad IHO ───────────────────────────────────────────────────
def tab_profundidad() -> None:
    st.header("📊 Análisis de Profundidad IHO S-5A")

    prof = cargar("profundidad")
    if prof is None:
        _aviso_script("analisis_profundidad.py")
        return

    resumen = prof.get("resumen", {})
    topicos = prof.get("topicos", [])
    por_asig = prof.get("por_asignatura", {})

    # KPIs
    c1, c2, c3 = st.columns(3)
    c1.metric("🔴 Sobredimensionados", resumen.get("sobredimensionados", 0))
    c2.metric("🟢 Alineados", resumen.get("alineados", 0))
    c3.metric("🔵 Subdimensionados", resumen.get("subdimensionados", 0))

    # Filtro por asignatura
    asigs_disp = sorted(por_asig.keys())
    asig_sel = st.multiselect(
        "Filtrar por asignatura (vacío = todas)",
        options=asigs_disp,
        default=[],
        placeholder="Todas las asignaturas",
    )

    top_filtrados = topicos
    if asig_sel:
        top_filtrados = [
            t for t in topicos
            if any(a in asig_sel for a in t.get("asignaturas_relacionadas", []))
        ]

    if top_filtrados:
        df_top = pd.DataFrame([
            {
                "Tópico": t["codigo_topico_iho"],
                "Descripción": (t["descripcion"][:55] + "…") if len(t["descripcion"]) > 55 else t["descripcion"],
                "Req. IHO": t["nivel_iho_requerido"],
                "Actual": t["nivel_actual_programa"],
                "Req. num": NIVEL_MAP.get(t["nivel_iho_requerido"], 2),
                "Act. num": NIVEL_MAP.get(t["nivel_actual_programa"], 2),
                "Clasificación": t["clasificacion"],
                "Asignaturas": ", ".join(t.get("asignaturas_relacionadas", [])),
            }
            for t in top_filtrados
        ])

        st.subheader("Nivel actual vs. requerido IHO por tópico")
        fig_bar = go.Figure()
        fig_bar.add_bar(
            x=df_top["Tópico"], y=df_top["Req. num"],
            name="Nivel requerido IHO", marker_color="#78909C",
        )
        fig_bar.add_bar(
            x=df_top["Tópico"], y=df_top["Act. num"],
            name="Nivel actual programa",
            marker_color=[CLASIF_COLOR.get(c, "#607D8B") for c in df_top["Clasificación"]],
        )
        fig_bar.update_layout(
            barmode="group",
            yaxis=dict(
                title="Nivel IHO S-5A",
                range=[0, 3.5],
                tickvals=[1, 2, 3],
                ticktext=["B – Básico", "I – Intermedio", "A – Avanzado"],
            ),
            height=400, legend=dict(orientation="h"),
        )
        st.plotly_chart(fig_bar, width="stretch")

        # Tabla + pie en columnas
        col_tab, col_pie = st.columns([3, 1])
        with col_tab:
            st.dataframe(
                df_top[["Tópico", "Descripción", "Req. IHO", "Actual", "Clasificación", "Asignaturas"]]
                .style.map(
                    lambda v: f"background-color:{CLASIF_COLOR.get(v,'white')};color:white",
                    subset=["Clasificación"],
                ),
                width="stretch", height=380,
            )
        with col_pie:
            conteo_c = df_top["Clasificación"].value_counts().reset_index()
            conteo_c.columns = ["Clasificación", "N"]
            fig_pie = px.pie(
                conteo_c, values="N", names="Clasificación",
                color="Clasificación",
                color_discrete_map=CLASIF_COLOR,
            )
            fig_pie.update_layout(height=320, showlegend=False, margin=dict(t=10, b=10))
            st.plotly_chart(fig_pie, width="stretch")

    # Resumen por asignatura
    st.subheader("Resumen por asignatura")
    if por_asig:
        rows_asig = [
            {
                "Asignatura": cod,
                "🔴 Sobredim.": v.get("SOBREDIMENSIONADO", 0),
                "🟢 Alineados": v.get("ALINEADO", 0),
                "🔵 Subdim.": v.get("SUBDIMENSIONADO", 0),
            }
            for cod, v in por_asig.items()
        ]
        df_asig = pd.DataFrame(rows_asig).sort_values("🔴 Sobredim.", ascending=False)
        st.dataframe(df_asig, width="stretch")


# ── Tab 4 — Propuesta Curricular ──────────────────────────────────────────────
def tab_propuesta() -> None:
    st.header("📋 Propuesta Curricular")

    prop = cargar("propuesta")
    if prop is None:
        _aviso_script("propuesta_curricular.py")
        return

    resumen = prop.get("resumen", {})
    comp = prop.get("comparativa_internacional", {})
    prop_asigs = prop.get("propuesta_asignaturas", [])

    # KPIs principales
    st.subheader("📈 Métricas de impacto")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Horas actuales", f"{resumen.get('total_horas_actuales', 0):,}")
    c2.metric(
        "Horas propuestas", f"{resumen.get('total_horas_propuestas', 0):,}",
        delta=f"−{resumen.get('reduccion_horas', 0):,} h",
        delta_color="inverse",
    )
    c3.metric("Reducción", f"{resumen.get('reduccion_porcentaje', 0):.1f}%")
    c4.metric("Cobertura IHO", f"{resumen.get('cobertura_iho_mantenida_porcentaje', 0):.1f}%")

    dist = resumen.get("distribucion_semestres_horas", {})
    d1, d2, d3 = st.columns(3)
    d1.metric("Semestre 1", f"{dist.get('semestre_1', 0):,} h")
    d2.metric("Semestre 2", f"{dist.get('semestre_2', 0):,} h")
    d3.metric("Semestre 3", f"{dist.get('semestre_3', 0):,} h")

    # Comparativa internacional
    st.subheader("🌍 Comparativa internacional")
    ci1, ci2, ci3 = st.columns(3)
    ci1.metric("SHOA actual (tópicos IHO)", f"{comp.get('referencia_horas_shoa_topics', 0):.0f} h")
    ci2.metric("Promedio internacional", f"{comp.get('promedio_internacional_topics', 0):.0f} h")
    brecha = comp.get("brecha_vs_promedio_internacional", 0)
    ci3.metric("Brecha propuesta vs. intl.", f"{brecha:+.0f} h")

    # ── Nueva malla curricular con dos versiones ─────────────────────────────
    st.markdown("---")
    st.subheader("📅 Nueva malla curricular — 3 semestres (1.5 años)")

    malla_va = prop.get("malla_v_a", [])
    malla_vb = prop.get("malla_v_b", [])

    if not malla_va:
        st.warning("Ejecuta `propuesta_curricular.py` para generar las versiones de malla.")
    else:
        version_sel = st.radio(
            "Versión de la malla:",
            ["Versión A — PH401 al final (estructura tradicional)",
             "Versión B — PH401 distribuido (exposición temprana al campo)"],
            horizontal=True,
        )
        malla_activa = malla_va if "Versión A" in version_sel else malla_vb

        sems_m: dict[int, list[dict]] = {1: [], 2: [], 3: []}
        for slot in malla_activa:
            s = slot.get("semestre", 1)
            if s in sems_m:
                sems_m[s].append(slot)

        cols_sem = st.columns(3)
        for sem, col in enumerate(cols_sem, 1):
            col.markdown(f"### Semestre {sem}")
            total_sem = sum(sl["horas_propuestas"]["total"] for sl in sems_m[sem])
            h_lect    = sum(sl["horas_propuestas"]["total"] for sl in sems_m[sem]
                           if "PH401" not in sl["codigo"])
            h_campo   = total_sem - h_lect
            detalle   = (f"**{total_sem:,} h** ✅" if h_campo == 0 and total_sem <= 500
                         else f"**{h_lect} h lectivas + {h_campo} h campo**")
            col.markdown(detalle, unsafe_allow_html=True)
            for slot in sems_m[sem]:
                color = ACCION_COLOR.get(slot["accion"], "#607D8B")
                hp    = slot["horas_propuestas"]
                nota_html = (f"<br><span style='font-size:10px;opacity:0.85'>{slot['nota']}</span>"
                             if slot.get("nota") else "")
                col.markdown(
                    f"<div style='background:{color};color:#fff;border-radius:6px;"
                    f"padding:8px 10px;margin-bottom:6px'>"
                    f"<b style='font-size:12px'>{slot['codigo']}</b><br>"
                    f"<span style='font-size:11px'>{slot['nombre']}</span><br>"
                    f"<span style='font-size:11px'>"
                    f"Total: {hp['total']} h | T:{hp['T']} P:{hp['P']} AE:{hp['SG']}"
                    f"</span>{nota_html}</div>",
                    unsafe_allow_html=True,
                )

        # Leyenda acciones
        st.markdown("**Leyenda de acciones:**")
        ley_cols = st.columns(3)
        for i, (accion, color) in enumerate(ACCION_COLOR.items()):
            ley_cols[i % 3].markdown(
                f"<div style='background:{color};color:#fff;border-radius:4px;"
                f"padding:3px 6px;text-align:center;font-size:11px;margin-bottom:4px'>{accion}</div>",
                unsafe_allow_html=True,
            )

        # ── Tabla comparativa A vs B ─────────────────────────────────────────
        st.markdown("---")
        st.subheader("📊 Comparativa Versión A vs. Versión B")

        def _sem_stats(malla: list[dict]) -> dict[int, dict]:
            from collections import defaultdict
            sd: dict = defaultdict(list)
            for s in malla:
                sd[s["semestre"]].append(s)
            return {
                sem: {
                    "total": sum(sl["horas_propuestas"]["total"] for sl in slots),
                    "lectivas": sum(sl["horas_propuestas"]["total"] for sl in slots
                                   if "PH401" not in sl["codigo"]),
                    "n_slots": len(slots),
                }
                for sem, slots in sd.items()
            }

        sa = _sem_stats(malla_va)
        sb = _sem_stats(malla_vb)

        def _fmt(stats: dict, sem: int) -> str:
            s = stats.get(sem, {})
            h = s.get("total", 0)
            lect = s.get("lectivas", 0)
            campo = h - lect
            return f"{h} h" + (f" ({lect}h lect. + {campo}h campo)" if campo else "")

        df_comp = pd.DataFrame([
            {"Criterio": "Horas Semestre 1 (Fundamentos)",
             "Versión A": _fmt(sa, 1), "Versión B": _fmt(sb, 1)},
            {"Criterio": "Horas Semestre 2 (Operaciones)",
             "Versión A": _fmt(sa, 2), "Versión B": _fmt(sb, 2)},
            {"Criterio": "Horas Semestre 3 (Procesamiento + Práctica)",
             "Versión A": _fmt(sa, 3), "Versión B": _fmt(sb, 3)},
            {"Criterio": "Exposición temprana al trabajo de campo",
             "Versión A": "No — PH401 inicia en Sem. 3",
             "Versión B": "Sí — salidas a terreno desde Sem. 2"},
            {"Criterio": "Riesgo académico",
             "Versión A": "Bajo — prereqs cubiertos antes del proyecto",
             "Versión B": "Medio — requiere coordinación logística en Sem. 2"},
            {"Criterio": "Flexibilidad operativa",
             "Versión A": "Baja — proyecto concentrado al final",
             "Versión B": "Alta — integración progresiva con el terreno"},
            {"Criterio": "Recomendación",
             "Versión A": "Estructura tradicional; primer ciclo completo antes del proyecto",
             "Versión B": "Exposición gradual; apropiada si hay recursos logísticos en Sem. 2"},
        ])
        st.dataframe(df_comp, width="stretch")

    # Tabla comparativa completa
    st.markdown("---")
    st.subheader("📊 Tabla comparativa detallada")
    if prop_asigs:
        df_prop = pd.DataFrame([
            {
                "Asignatura":  p["asignatura"],
                "Nombre":      p["nombre_asignatura"],
                "Acción":      p["accion"],
                "T act.":      p["horas_actuales"]["T"],
                "P act.":      p["horas_actuales"]["P"],
                "AE act.":     p["horas_actuales"]["SG"],
                "Total act.":  p["horas_actuales"]["total"],
                "T prop.":     p["horas_propuestas"]["T"],
                "P prop.":     p["horas_propuestas"]["P"],
                "AE prop.":    p["horas_propuestas"]["SG"],
                "Total prop.": p["horas_propuestas"]["total"],
                "Sem. act.":   p["semestre_actual"],
                "Sem. prop.":  p["semestre_propuesto"],
                "Justificación": p["justificacion"],
            }
            for p in prop_asigs
        ])
        st.dataframe(
            df_prop.style.map(
                lambda v: f"background-color:{ACCION_COLOR.get(v,'#607D8B')};color:white",
                subset=["Acción"],
            ),
            width="stretch",
            height=480,
        )

    # Gráfico horas actuales vs propuestas
    if prop_asigs:
        df_bar = pd.DataFrame([
            {
                "Asignatura": p["asignatura"],
                "Horas actuales": p["horas_actuales"]["total"],
                "Horas propuestas": p["horas_propuestas"]["total"],
            }
            for p in sorted(prop_asigs, key=lambda x: x["semestre_propuesto"])
        ])
        fig_comp = px.bar(
            df_bar.melt(id_vars="Asignatura", var_name="Versión", value_name="Horas"),
            x="Asignatura", y="Horas", color="Versión", barmode="group",
            color_discrete_map={"Horas actuales": "#546E7A", "Horas propuestas": "#1E88E5"},
            title="Comparación horas actuales vs. propuestas por asignatura",
        )
        fig_comp.update_layout(height=380)
        st.plotly_chart(fig_comp, width="stretch")


# ── Tab 5 — Exportar Word ────────────────────────────────────────────────────
def _generar_word(prog: dict, reit: dict, prof: dict, prop: dict) -> bytes:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = DocxDocument()

    def _h(text: str, level: int = 1) -> None:
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _p(text: str) -> None:
        doc.add_paragraph(text)

    def _tabla(filas: list[list[str]]) -> None:
        if not filas:
            return
        t = doc.add_table(rows=len(filas), cols=len(filas[0]))
        t.style = "Table Grid"
        for i, fila in enumerate(filas):
            for j, val in enumerate(fila):
                cell = t.rows[i].cells[j]
                cell.text = str(val)
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

    # Portada
    if LOGO_PATH.exists():
        try:
            doc.add_picture(str(LOGO_PATH), width=Inches(2.0))
        except Exception:
            pass

    par_titulo = doc.add_paragraph()
    par_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = par_titulo.add_run("PROPUESTA DE REDISEÑO CURRICULAR\nPrograma de Hidrografía — SHOA\n")
    run_t.bold = True
    run_t.font.size = Pt(18)

    par_sub = doc.add_paragraph()
    par_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par_sub.add_run(
        f"Ajuste a 1 año y medio (3 semestres)\n"
        f"Estándares IHO S-5A\n"
        f"Generado: {date.today().strftime('%d/%m/%Y')}"
    )
    doc.add_page_break()

    # 1. Resumen ejecutivo
    _h("1. Resumen Ejecutivo")
    res = prop.get("resumen", {})
    _p(
        "Este documento presenta los resultados del análisis curricular del programa de "
        "Hidrografía de la Armada de Chile (SHOA) y propone su reestructuración a 3 semestres "
        "(1 año y medio), alineado con los estándares IHO S-5A."
    )
    _tabla([
        ["Indicador", "Valor"],
        ["Horas actuales", str(res.get("total_horas_actuales", "—"))],
        ["Horas propuestas", str(res.get("total_horas_propuestas", "—"))],
        ["Reducción de horas", str(res.get("reduccion_horas", "—"))],
        ["Reducción porcentual", f"{res.get('reduccion_porcentaje', 0):.1f}%"],
        ["Cobertura IHO mantenida", f"{res.get('cobertura_iho_mantenida_porcentaje', 0):.1f}%"],
        ["Cabe en 1.5 años", "Sí" if res.get("curriculo_cabe_en_1_5_anios") else "No"],
    ])

    # 2. Análisis de progresión
    _h("2. Análisis de Progresión")
    r_prog = prog.get("resumen", {})
    _p(
        f"Asignaturas: {r_prog.get('total_asignaturas', 0)} | "
        f"Dependencias: {r_prog.get('total_dependencias', 0)} | "
        f"Problemas detectados: {r_prog.get('total_problemas', 0)}"
    )
    problemas_prog = prog.get("problemas_progresion", [])
    if problemas_prog:
        _h("2.1 Problemas de Progresión", level=2)
        filas_pb = [["#", "Tipo", "Asignatura", "Sem.", "Descripción"]]
        for i, pb in enumerate(problemas_prog, 1):
            desc = pb.get("descripcion", "")
            filas_pb.append([
                str(i),
                pb.get("tipo", "").replace("_", " ").title(),
                pb.get("asignatura", "—"),
                str(pb.get("semestre_asignatura", "—")),
                desc[:110] + ("…" if len(desc) > 110 else ""),
            ])
        _tabla(filas_pb)

    sin_prereq = prog.get("asignaturas_sin_prerrequisitos_claros", [])
    if sin_prereq:
        _h("2.2 Asignaturas sin Prerrequisitos Claros", level=2)
        _tabla(
            [["Código", "Nombre", "Nivel IHO"]]
            + [[a["codigo"], a.get("nombre", ""), str(a.get("nivel_iho_predominante", ""))] for a in sin_prereq]
        )

    # 3. Matriz de reiteración
    _h("3. Matriz de Reiteración de Contenidos")
    r_reit = reit.get("resumen", {})
    _p(
        f"Pares de contenidos similares (umbral ≥{r_reit.get('threshold_fuzzy', 75)}%): "
        f"{r_reit.get('total_reiteraciones_detectadas', 0)} | "
        f"Innecesarias: {r_reit.get('total_innecesarias', 0)} | "
        f"Horas recuperables: {r_reit.get('horas_recuperables_estimadas', 0):.1f} h"
    )
    innecesarias = [r for r in reit.get("reiteraciones", []) if r["tipo_reiteracion"] == "INNECESARIA"]
    if innecesarias:
        _h("3.1 Reiteraciones Innecesarias (muestra)", level=2)
        filas_ri = [["Contenido (recortado)", "Asig. 1", "Asig. 2", "Similitud", "Horas dup."]]
        for r in innecesarias[:20]:
            filas_ri.append([
                r["contenido"][:80],
                r["asignatura_1"], r["asignatura_2"],
                f"{r['similitud']}%", str(r["horas_duplicadas_estimadas"]),
            ])
        _tabla(filas_ri)

    # 4. Profundidad IHO
    _h("4. Análisis de Profundidad IHO S-5A")
    r_prof = prof.get("resumen", {})
    _p(
        f"Tópicos analizados: {r_prof.get('total_topicos_analizados', 0)} | "
        f"Sobredimensionados: {r_prof.get('sobredimensionados', 0)} | "
        f"Alineados: {r_prof.get('alineados', 0)} | "
        f"Subdimensionados: {r_prof.get('subdimensionados', 0)}"
    )
    por_asig = prof.get("por_asignatura", {})
    if por_asig:
        _h("4.1 Resumen por Asignatura", level=2)
        _tabla(
            [["Asignatura", "Sobredim.", "Alineados", "Subdim."]]
            + [[cod, str(v.get("SOBREDIMENSIONADO", 0)), str(v.get("ALINEADO", 0)), str(v.get("SUBDIMENSIONADO", 0))]
               for cod, v in sorted(por_asig.items())]
        )

    # 5. Propuesta de nueva estructura
    _h("5. Propuesta de Nueva Estructura Curricular")
    _p(
        "El currículo se organiza en 3 semestres siguiendo la progresión IHO: "
        "Fundamentos (Sem.1) → Operaciones (Sem.2) → Procesamiento y Gestión (Sem.3)."
    )
    prop_asigs = prop.get("propuesta_asignaturas", [])
    if prop_asigs:
        filas_prop = [[
            "Asignatura", "Acción",
            "T act.", "P act.", "AE act.", "Total act.",
            "T prop.", "P prop.", "AE prop.", "Total prop.",
            "Sem. act.", "Sem. prop.", "Justificación",
        ]]
        for pa in prop_asigs:
            ha, hp = pa["horas_actuales"], pa["horas_propuestas"]
            filas_prop.append([
                pa["asignatura"], pa["accion"],
                str(ha["T"]), str(ha["P"]), str(ha["SG"]), str(ha["total"]),
                str(hp["T"]), str(hp["P"]), str(hp["SG"]), str(hp["total"]),
                str(pa["semestre_actual"]), str(pa["semestre_propuesto"]),
                pa["justificacion"][:100],
            ])
        _tabla(filas_prop)

    # 6. Métricas de impacto
    _h("6. Métricas de Impacto")
    comp = prop.get("comparativa_internacional", {})
    dist = res.get("distribucion_semestres_horas", {})
    _tabla([
        ["Métrica", "Valor"],
        ["Horas actuales", str(res.get("total_horas_actuales", "—"))],
        ["Horas propuestas", str(res.get("total_horas_propuestas", "—"))],
        ["Reducción", f"{res.get('reduccion_porcentaje', 0):.1f}%"],
        ["Cobertura IHO S-5A", f"{res.get('cobertura_iho_mantenida_porcentaje', 0):.1f}%"],
        ["Horas semestre 1", str(dist.get("semestre_1", "—"))],
        ["Horas semestre 2", str(dist.get("semestre_2", "—"))],
        ["Horas semestre 3", str(dist.get("semestre_3", "—"))],
        ["Promedio internacional (tópicos)", f"{comp.get('promedio_internacional_topics', 0):.0f} h"],
        ["Brecha propuesta vs. intl.", f"{comp.get('brecha_vs_promedio_internacional', 0):+.0f} h"],
        ["Reiteraciones innecesarias", str(r_reit.get("total_innecesarias", "—"))],
        ["Horas recuperadas (reiteración)", f"{r_reit.get('horas_recuperables_estimadas', 0):.1f} h"],
    ])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def tab_exportar() -> None:
    st.header("📄 Exportar Propuesta en Word")
    st.info(
        "Genera un documento Word formal con portada institucional, "
        "todos los análisis y la propuesta de nueva malla curricular."
    )

    requeridos = {
        "progresion":  "analisis_progresion.py",
        "reiteracion": "matriz_reiteracion.py",
        "profundidad": "analisis_profundidad.py",
        "propuesta":   "propuesta_curricular.py",
    }
    datos: dict[str, dict] = {}
    faltantes: list[str] = []
    for key, script in requeridos.items():
        d = _leer(str(PATHS[key]))
        if d is None:
            faltantes.append(f"`{PATHS[key].name}` → ejecuta `python {script}`")
        else:
            datos[key] = d

    if LOGO_PATH.exists():
        st.success("✅ Logo SHOA encontrado en `assets/logo_shoa.png`")
    else:
        st.info("ℹ️ No se encontró `assets/logo_shoa.png`. El documento se generará sin logo.")

    if faltantes:
        st.warning("⚠️ Faltan archivos:\n\n" + "\n".join(f"- {f}" for f in faltantes))
        st.button("📄 Generar Propuesta Curricular en Word", disabled=True)
        return

    if st.button("📄 Generar Propuesta Curricular en Word", type="primary"):
        with st.spinner("Generando documento Word…"):
            try:
                doc_bytes = _generar_word(
                    prog=datos["progresion"],
                    reit=datos["reiteracion"],
                    prof=datos["profundidad"],
                    prop=datos["propuesta"],
                )
                st.download_button(
                    label="⬇️ Descargar Propuesta_Curricular_SHOA.docx",
                    data=doc_bytes,
                    file_name="Propuesta_Curricular_SHOA.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                st.success("✅ Documento generado correctamente.")
            except Exception as exc:
                st.error(f"Error generando documento: {exc}")


# ── Tab 6 — Validación de Datos ──────────────────────────────────────────────
def tab_validacion() -> None:
    st.header("🔍 Validación de Datos — Programas Estructurados")
    st.caption(
        "Contenido extraído de `programas_estructurados.json`. "
        "Compara visualmente contra el Word original sin abrir el JSON."
    )

    prog_data = cargar("programas", mostrar_aviso=True)
    curr_data = cargar("curriculum", mostrar_aviso=False)
    if prog_data is None:
        _aviso_script("extractor_programas.py")
        return

    programas = [p for p in prog_data.get("programas", []) if p.get("codigo_asignatura")]
    if not programas:
        st.warning("No se encontraron programas válidos.")
        return

    # Horas corregidas desde curriculum_data.json (fuente autoritativa)
    import re as _re_val
    def _nc(s: str) -> str:
        return _re_val.sub(r"[^a-z0-9]", "", s.lower())

    horas_curr: dict[str, dict] = {}
    if curr_data:
        for cod_raw, v in curr_data.get("horas_asignaturas", {}).items():
            horas_curr[_nc(cod_raw)] = v

    # Selector de asignatura
    opciones = {f"{p['codigo_asignatura']} — {p.get('nombre_asignatura', '')}": p for p in programas}
    sel_label = st.selectbox("Seleccionar asignatura:", list(opciones.keys()))
    prog = opciones[sel_label]
    cod  = prog["codigo_asignatura"]

    st.markdown("---")

    # ── Cabecera ──────────────────────────────────────────────────────────────
    col_head, col_sem = st.columns([4, 1])
    col_head.subheader(f"{cod} — {prog.get('nombre_asignatura', '')}")
    sem = prog.get("semestre")
    col_sem.metric("Semestre actual", sem if sem else "—")

    if prog.get("error"):
        st.error(f"Error en extracción: {prog['error']}")
        return

    # ── Horas: curriculum_data.json (Excel) como fuente primaria ─────────────
    # El campo horas en programas_estructurados.json extrae el TOTAL del Word
    # para T, P y AE (bug estructural de tabla horizontal). curriculum_data es
    # la fuente autoritativa (Excel Tabla Resumen).
    horas_doc  = prog.get("horas", {})
    horas_real = horas_curr.get(_nc(cod), {}) if curr_data else {}

    # curriculum_data usa "Total" (mayúscula), metricas usan "total" (minúscula)
    def _get_h(d: dict, key: str) -> int:
        return int(d.get(key) or d.get(key.capitalize()) or d.get(key.upper()) or 0)

    # Detectar si el Word tiene el bug T==P==SG==total (extracción errónea)
    _t  = _get_h(horas_doc, "T")
    _p  = _get_h(horas_doc, "P")
    _sg = _get_h(horas_doc, "SG")
    _tot= _get_h(horas_doc, "total")
    _word_bugged = (_t == _p == _sg == _tot and _t > 0)

    st.subheader("Horas")
    if not horas_real:
        st.caption("⚠️ Sin datos en curriculum_data.json para esta asignatura.")
    c1, c2, c3, c4 = st.columns(4)
    metricas = [
        ("Teoría (T)",        "T",     "#1565C0"),
        ("Práctica (P)",      "P",     "#2E7D32"),
        ("Auto Estudio (AE)", "SG",    "#E65100"),
        ("Total",             "total", "#6A1B9A"),
    ]
    for col, (label, key, color) in zip([c1, c2, c3, c4], metricas):
        val_excel = _get_h(horas_real, key) if horas_real else None
        val_word  = _get_h(horas_doc, key)
        # Primario: Excel (fuente autoritativa). Fallback: Word si no hay Excel.
        val_show  = val_excel if val_excel is not None else val_word
        fuente    = "Excel Tabla Resumen" if val_excel is not None else "Word extraído"
        col.markdown(
            f"<div style='background:{color};color:#fff;border-radius:8px;"
            f"padding:12px;text-align:center'>"
            f"<div style='font-size:11px;opacity:0.85'>{label}</div>"
            f"<div style='font-size:26px;font-weight:bold'>{val_show}</div>"
            f"<div style='font-size:10px;opacity:0.75'>{fuente}</div>"
            f"</div>",
            unsafe_allow_html=True,
        )
        # Nota si Word (fila TOTAL) tiene valor distinto al Excel
        if val_excel is not None and val_word != val_excel:
            if _word_bugged:
                col.caption("⚠️ Word: extracción incorrecta (tabla horizontal)")
            elif val_word > 0:
                col.caption(f"ℹ️ Word TOTAL: {val_word} h")

    # ── Prerrequisitos ─────────────────────────────────────────────────────────
    prereqs = prog.get("prerrequisitos", [])
    st.subheader("Prerrequisitos")
    if prereqs:
        st.markdown(" | ".join(f"`{r}`" for r in prereqs))
    else:
        st.caption("Sin prerrequisitos declarados.")

    # ── Resultados de aprendizaje ──────────────────────────────────────────────
    ra_list = prog.get("resultados_aprendizaje", [])
    st.subheader(f"Resultados de Aprendizaje ({len(ra_list)})")
    if ra_list:
        for i, ra in enumerate(ra_list, 1):
            st.markdown(
                f"<div style='border-left:3px solid #1565C0;padding:6px 12px;"
                f"margin-bottom:6px;background:#f5f8ff;border-radius:0 4px 4px 0'>"
                f"<b style='font-size:11px;color:#1565C0'>RA {i}</b><br>"
                f"<span style='font-size:13px'>{ra}</span></div>",
                unsafe_allow_html=True,
            )
    else:
        st.caption("Sin resultados de aprendizaje extraídos.")

    # ── Unidades temáticas ─────────────────────────────────────────────────────
    unidades = prog.get("unidades_tematicas", [])
    st.subheader(f"Unidades Temáticas ({len(unidades)})")

    if not unidades:
        st.caption("Sin unidades temáticas extraídas.")
    else:
        total_h_unidades = sum(int(u.get("horas_asignadas") or 0) for u in unidades)
        st.caption(f"Total horas declaradas en unidades: **{total_h_unidades} h**")

        for idx, u in enumerate(unidades):
            nombre_u = u.get("nombre_unidad", f"Unidad {idx + 1}")
            horas_u  = u.get("horas_asignadas")
            contenidos   = u.get("contenidos_especificos", [])
            actividades  = u.get("actividades_asociadas", [])
            h_label = f"— {horas_u} h" if horas_u else ""

            with st.expander(f"**{nombre_u}** {h_label}  ({len(contenidos)} contenidos)"):
                # Barra de horas de la unidad
                if horas_u and total_h_unidades:
                    pct = int(horas_u) / total_h_unidades * 100
                    st.markdown(
                        f"<div style='background:#e0e0e0;border-radius:4px;height:8px;margin-bottom:8px'>"
                        f"<div style='background:#1565C0;width:{pct:.0f}%;height:8px;border-radius:4px'>"
                        f"</div></div>",
                        unsafe_allow_html=True,
                    )
                    st.caption(f"{pct:.0f}% del total de horas de la asignatura")

                # Contenidos
                if contenidos:
                    st.markdown("**Contenidos específicos:**")
                    for j, c in enumerate(contenidos, 1):
                        st.markdown(
                            f"<div style='padding:3px 0 3px 12px;border-left:2px solid #90CAF9;"
                            f"margin-bottom:3px;font-size:13px'>"
                            f"<span style='color:#1565C0;font-weight:bold'>{j}.</span> {c}</div>",
                            unsafe_allow_html=True,
                        )

                # Actividades
                if actividades:
                    st.markdown("**Actividades asociadas:**")
                    for k, act in enumerate(actividades, 1):
                        # Truncar texto muy largo
                        texto = act if len(act) <= 300 else act[:300] + "…"
                        st.markdown(
                            f"<div style='padding:3px 0 3px 12px;border-left:2px solid #A5D6A7;"
                            f"margin-bottom:3px;font-size:12px;color:#444'>"
                            f"<span style='color:#2E7D32;font-weight:bold'>{k}.</span> {texto}</div>",
                            unsafe_allow_html=True,
                        )

    # ── Campos no encontrados ──────────────────────────────────────────────────
    campos_faltantes = prog.get("campos_no_encontrados", [])
    if campos_faltantes:
        st.markdown("---")
        st.caption(f"Campos no extraídos del Word: {', '.join(campos_faltantes)}")


# ── Main ─────────────────────────────────────────────────────────────────────
def main() -> None:
    st.set_page_config(
        page_title="Diseño Curricular SHOA",
        page_icon="🧭",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    # ── Sidebar ───────────────────────────────────────────────────────────────
    with st.sidebar:
        if LOGO_PATH.exists():
            st.image(str(LOGO_PATH), width=120)
        else:
            st.markdown("## 🧭")
        st.markdown("### Diseño Curricular SHOA")
        st.markdown("*Reestructuración IHO S-5A*")
        st.divider()
        st.markdown(
            f"📊 **Datos:** Versión {_DATOS_VERSION}  \n"
            "🌐 **Modo:** Online"
        )
        st.divider()
        st.caption("Tabs disponibles:")
        st.markdown(
            "🗺️ Mapa de Progresión  \n"
            "🔁 Matriz de Reiteración  \n"
            "📊 Profundidad IHO  \n"
            "📋 Propuesta Curricular  \n"
            "📄 Exportar Word  \n"
            "🔍 Validación de Datos"
        )
        st.divider()
        st.caption("Desarrollado para SHOA · Chile")

    col_logo, col_tit = st.columns([1, 7])
    with col_logo:
        if LOGO_PATH.exists():
            st.image(str(LOGO_PATH), width=80)
        else:
            st.markdown("## 🧭")
    with col_tit:
        st.title("Diseño Curricular SHOA — Fase 2")
        st.markdown(
            "*Reestructuración del programa de Hidrografía a 1.5 años según estándares IHO S-5A*"
        )

    st.markdown("---")

    t1, t2, t3, t4, t5, t6 = st.tabs([
        "🗺️ Mapa de Progresión",
        "🔁 Matriz de Reiteración",
        "📊 Profundidad IHO",
        "📋 Propuesta Curricular",
        "📄 Exportar Word",
        "🔍 Validación de Datos",
    ])
    with t1:
        tab_progresion()
    with t2:
        tab_reiteracion()
    with t3:
        tab_profundidad()
    with t4:
        tab_propuesta()
    with t5:
        tab_exportar()
    with t6:
        tab_validacion()


if __name__ == "__main__":
    main()
